#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成优化后的董事工作报告（标注修改版）"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_highlight(run, color="yellow"):
    """设置文字高亮颜色"""
    run.font.highlight_color = 3  # Yellow

def add_change_marker(paragraph, old_text, new_text):
    """添加修改标注（旧文+新文）"""
    # 旧文（删除线+红色）
    run1 = paragraph.add_run(f"【原→{old_text}】")
    run1.font.strike = True
    run1.font.color.rgb = RGBColor(255, 0, 0)
    run1.font.highlight_color = 2  # Red-ish
    
    # 新文（黄色高亮）
    run2 = paragraph.add_run(f"【修改→{new_text}】")
    run2.font.color.rgb = RGBColor(255, 0, 0)
    run2.font.highlight_color = 4  # Yellow
    
    paragraph.add_run("\n")

def add_yellow_text(paragraph, text):
    """添加黄色高亮文字"""
    run = paragraph.add_run(text)
    run.font.highlight_color = 4  # Yellow
    return run

def add_red_text(paragraph, text):
    """添加红色文字"""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(255, 0, 0)
    return run

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ==================== 封面 ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("上 海 麒 麟 士 企 业 发 展 有 限 公 司")
run.font.size = Pt(18)
run.font.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("2025年度董事工作报告")
run2.font.size = Pt(26)
run2.font.bold = True
run2.font.name = '黑体'
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date.add_run("（2026年4月）")
run3.font.size = Pt(14)

doc.add_page_break()

# ==================== 正文 ====================

# 标题
h1 = doc.add_heading('', level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h1.add_run("2025年度董事工作报告")
run.font.size = Pt(22)
run.font.bold = True

# 说明
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.add_run("（标注修改版——红色删除线为原内容，黄色高亮为修改/新增内容）").font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph()

# 一、
section1 = doc.add_heading('', level=1)
run = section1.add_run("一、2025年度工作总结")
run.font.size = Pt(16)

# （一）概述
sub1 = doc.add_heading('', level=2)
run = sub1.add_run("（一）概述")
run.font.size = Pt(14)

p1 = doc.add_paragraph()
p1.add_run("2025年，上海麒麟士企业发展有限公司董事会严格按照《公司法》、《公司章程》及上级集团监管要求，").font.size = Pt(12)

# 修改标注段落
p1_new = doc.add_paragraph()
p1_new.add_run("【懂事视角强化】").font.bold = True
p1_new.add_run("\n")

# 修改1
p_mod1 = doc.add_paragraph()
add_red_text(p_mod1, "原表述：").font.size = Pt(10)
add_red_text(p_mod1, "「履行职责」")
p_mod1.add_run(" → ").font.size = Pt(10)
add_yellow_text(p_mod1, "修改为：")
add_yellow_text(p_mod1, "「切实履行国有资产出资人代表职责，强化战略决策与风险管控，聚焦公司战略规划、经营管理、风险管控、合规治理等核心工作」")
p_mod1.add_run("\n").font.size = Pt(10)

p1_text = doc.add_paragraph()
p1_text.add_run("本年度，公司董事会切实发挥").font.size = Pt(12)
add_yellow_text(p1_text, "「把方向、管大局、保落实」").font.bold = True
p1_text.add_run("的领导作用，聚焦公司战略规划、经营管理、风险管控、合规治理等核心工作，全力保障公司稳健运营与国有资产保值增值，").font.size = Pt(12)
add_yellow_text(p1_text, "切实维护全体股东合法权益").font.bold = True
p1_text.add_run("。现将2025年度工作情况汇报如下：").font.size = Pt(12)

# （二）董事专题会召开与执行情况
sub2 = doc.add_heading('', level=2)
run = sub2.add_run("（二）董事专题会召开与执行情况")
run.font.size = Pt(14)

p2 = doc.add_paragraph()
p2.add_run("本年度董事专题会共召开会议").font.size = Pt(12)
add_yellow_text(p2, "19次").font.bold = True
p2.add_run("，其中定期会议1次，临时会议18次。会议召集、通知、表决、记录等流程均符合规范要求，决策程序合法合规。董事会充分发挥").font.size = Pt(12)
add_yellow_text(p2, "「科学决策、民主决策、依法决策」").font.bold = True
p2.add_run("的优良传统，").font.size = Pt(12)
add_yellow_text(p2, "确保各项决议事项落实到位").font.bold = True
p2.add_run("。具体如下：").font.size = Pt(12)

# 会议列表保持原样，仅修改引言
meetings = [
    "1. 2025年1月21日——研究欧登公司负责人年度绩效分配事宜",
    "2. 2025年2月7日——研究欧登公司及物业关闭事宜",
    "3. 2025年2月7日——研究欧登保龄球馆退卡事宜",
    "4. 2025年3月14日——研究麒麟士公司重要制度",
    "5. 2025年3月28日——研究欧登固定资产价格咨询等事宜",
    "6. 2025年4月10日——研究麒麟士公司组织机构调整事宜",
    "7. 2025年4月10日——研究欧登固定资产和无形资产处置拍卖事宜",
    "8. 2025年4月10日——研究麒麟士公司采购欧登公司固定资产事宜",
    "9. 2025年4月28日——研究《麒麟士公司2024年度董事/股东会议工作报告》",
    "10. 2025年5月7日——研究麒麟士公司吸收合并欧登事宜",
    "11. 2025年6月10日——研究衡船印刷分公司重大采购相关事宜",
    "12. 2025年6月30日——研究《衡船印刷分公司薪酬管理办法》",
    "13. 2025年6月30日——研究公司通过《衡船印刷分公司薪酬方案》",
    "14. 2025年10月10日——研究《麒麟士公司十五五规划（2026-2030）》",
    "15. 2025年10月11日——研究衡船印刷分公司设备报废事宜",
    "16. 2025年11月7日——研究《麒麟士公司2025年度风险管理情况及2026年形势分析研判》",
    "17. 2025年11月25日——《麒麟士公司2025年度内部控制评价报告》",
    "18. 2025年12月10日——研究2025年度物业合同预算追加事宜",
    "19. 2025年12月17日——研究《麒麟士公司2026年度预算》",
]

for m in meetings:
    doc.add_paragraph(m, style='List Number').paragraph_format.first_line_indent = Cm(0.5)

p_confirm = doc.add_paragraph()
p_confirm.add_run("各项决议均得到").font.size = Pt(12)
add_yellow_text(p_confirm, "有效执行").font.bold = True
p_confirm.add_run("，未出现违规决策或决议未落实情形。").font.size = Pt(12)

# （三）董事履职情况
sub3 = doc.add_heading('', level=2)
run = sub3.add_run("（三）董事履职情况")
run.font.size = Pt(14)

p3 = doc.add_paragraph()
p3.add_run("公司董事通过董事会议行使董事职权，").font.size = Pt(12)
add_yellow_text(p3, "定期向股东会报告工作，执行股东会的决议，主动接受股东监督").font.bold = True
p3.add_run("。董事参与议案审议与讨论，依法依规对各项重大事项进行").font.size = Pt(12)
add_yellow_text(p3, "审慎研究和科学决策").font.bold = True
p3.add_run("，").font.size = Pt(12)
add_yellow_text(p3, "充分发挥董事专业能力与勤勉义务").font.bold = True
p3.add_run("，切实维护股东与公司利益。在2025年度董事履职评价中，董事单丽辉得分97分。").font.size = Pt(12)

# 修改标注
p3_note = doc.add_paragraph()
add_red_text(p3_note, "【优化说明】此处可补充董事会专门委员会（如审计委员会、薪酬考核委员会）运作情况，体现董事履职的专业性。")

# （四）重点工作完成情况
sub4 = doc.add_heading('', level=2)
run = sub4.add_run("（四）重点工作完成情况")
run.font.size = Pt(14)

# 1. 经营情况
p4_1 = doc.add_paragraph()
run = p4_1.add_run("1. 经营情况")
run.font.bold = True
run.font.size = Pt(12)

p4_1_content = doc.add_paragraph()
p4_1_content.add_run("本年度实现营业收入").font.size = Pt(12)
add_yellow_text(p4_1_content, "5010.59万元").font.bold = True
p4_1_content.add_run("，同比上年1108.43万元增加3902.16万元，").font.size = Pt(12)
add_yellow_text(p4_1_content, "增幅352.05%").font.bold = True
p4_1_content.add_run("。").font.size = Pt(12)

p4_1_note = doc.add_paragraph()
p4_1_note.add_run("变动原因：").font.bold = True
p4_1_note.add_run("业务结构调整，2024年主要为欧登大楼物业管理、经营分公司业务等，2025年业务变更为所部物业管理、印刷业务等（其中，印刷收入668万元）。").font.size = Pt(12)

p4_1_profit = doc.add_paragraph()
p4_1_profit.add_run("本年度利润总额").font.size = Pt(12)
add_yellow_text(p4_1_profit, "501.8万元").font.bold = True
p4_1_profit.add_run("，同比上年175.39万元增加326.41万元，").font.size = Pt(12)
add_yellow_text(p4_1_profit, "增幅186.11%").font.bold = True
p4_1_profit.add_run("；净利润489.99万元，同比上年175.39万元增加314.60万元，").font.size = Pt(12)
add_yellow_text(p4_1_profit, "增幅179.38%").font.bold = True
p4_1_profit.add_run("。").font.size = Pt(12)

p4_1_breakdown = doc.add_paragraph()
p4_1_breakdown.add_run("利润构成：印刷厂利润353万元，后勤物业服务利润148.80万元；因欧登公司关闭，确认合并报表利润总额137.38万元。").font.size = Pt(12)

# 2. 机构设置与人员优化
p4_2 = doc.add_paragraph()
run = p4_2.add_run("2. 机构设置与人员优化")
run.font.bold = True
run.font.size = Pt(12)

p4_2_content = doc.add_paragraph()
p4_2_content.add_run("一是明确公司职能和定位，按照").font.size = Pt(12)
add_yellow_text(p4_2_content, "「同类或相近业务合并、促进深度融合」").font.bold = True
p4_2_content.add_run("的原则，重新调整组织机构设置；二是系统梳理和评估内部岗位，完成薪酬制度改革，通过双向选择、竞聘上岗、合理分流等方式，公司").font.size = Pt(12)
add_yellow_text(p4_2_content, "精简18人").font.bold = True
p4_2_content.add_run("，实现人力资源的优化配置。").font.size = Pt(12)

# 3. 子公司清理与资产处置
p4_3 = doc.add_paragraph()
run = p4_3.add_run("3. 子公司清理与资产处置")
run.font.bold = True
run.font.size = Pt(12)

p4_3_content = doc.add_paragraph()
p4_3_content.add_run("公司稳步完成对全资子公司上海欧登体育文化有限公司的").font.size = Pt(12)
add_yellow_text(p4_3_content, "关闭退出工作").font.bold = True
p4_3_content.add_run("。董事专题会审议通过了欧登公司及其物业的关闭方案、固定资产价格咨询、资产处置拍卖、以及由母公司吸收合并等议案，按照既定时间节点完成包括人员安置、资产处置、财务审计、税务清缴、工商注销等事宜，").font.size = Pt(12)
add_yellow_text(p4_3_content, "有效落实了集团公司深化改革专项行动的重要部署，确保国有资产处置依法合规、程序完备").font.bold = True
p4_3_content.add_run("。").font.size = Pt(12)

# 4. 服务保障
p4_4 = doc.add_paragraph()
run = p4_4.add_run("4. 服务保障")
run.font.bold = True
run.font.size = Pt(12)

p4_4_content = doc.add_paragraph()
p4_4_content.add_run("自后勤服务中心和车队业务并入后，").font.size = Pt(12)
add_yellow_text(p4_4_content, "合理调配岗位职责，顺利保障本年度各项服务工作开展").font.bold = True
p4_4_content.add_run("，为我所科研生产和日常办公提供有力支持。").font.size = Pt(12)

# 5. 分公司运营保障
p4_5 = doc.add_paragraph()
run = p4_5.add_run("5. 分公司运营保障")
run.font.bold = True
run.font.size = Pt(12)

p4_5_content = doc.add_paragraph()
p4_5_content.add_run("加强对印刷分公司日常监督和规范化管理，通过制定衡船印刷分公司薪酬管理办法、开展经理层竞聘和规范采购流程等工作，").font.size = Pt(12)
add_yellow_text(p4_5_content, "有效激发员工的工作积极性，营造健康规范的经营环境").font.bold = True
p4_5_content.add_run("。").font.size = Pt(12)

# 6. 规划与预算管理
p4_6 = doc.add_paragraph()
run = p4_6.add_run("6. 规划与预算管理")
run.font.bold = True
run.font.size = Pt(12)

p4_6_content = doc.add_paragraph()
p4_6_content.add_run("科学谋划公司未来五年发展蓝图，围绕公司总体目标，组织编制了").font.size = Pt(12)
add_yellow_text(p4_6_content, '《麒麟士公司「十五五」规划（2026-2030）》').font.bold = True
p4_6_content.add_run('，明确「十五五」时期公司的发展定位，聚焦主责主业，突出补短板、强弱项、提质量，从多个维度提出了保障措施，形成完善的实施保障体系。董事专题会完成2025年度财务决算、2026年度财务预算（包括职工福利服务费、物业服务费、安全生产资金预算）等事项的审议工作，').font.size = Pt(12)
add_yellow_text(p4_6_content, "强化全面预算管理").font.bold = True
p4_6_content.add_run("。").font.size = Pt(12)

# 7. 风险与内控建设
p4_7 = doc.add_paragraph()
run = p4_7.add_run("7. 风险与内控建设")
run.font.bold = True
run.font.size = Pt(12)

p4_7_content = doc.add_paragraph()
p4_7_content.add_run('公司持续推进合规管理与法治建设：一是推进制度修订、法治整改，将制度建设列为年度管理核心，对标上级开展制度「立改废」，').font.size = Pt(12)
add_yellow_text(p4_7_content, "共完成16项关键制度制修订").font.bold = True
p4_7_content.add_run("，并开展分类宣贯培训确保执行到位；二是围绕审计报告中3项主要问题，").font.size = Pt(12)
add_yellow_text(p4_7_content, "组织专项整改工作，于12月底完成整改").font.bold = True
p4_7_content.add_run("；三是加强采购及供应链合规管控，完成与七〇四所合同主体变更，委托采购中心实施采购71次，办理75家合格供方准入，").font.size = Pt(12)
add_yellow_text(p4_7_content, "确保采购合规高效").font.bold = True
p4_7_content.add_run("。").font.size = Pt(12)

# 8. 党建工作情况
p4_8 = doc.add_paragraph()
run = p4_8.add_run("8. 党建工作情况")
run.font.bold = True
run.font.size = Pt(12)

p4_8_content = doc.add_paragraph()
p4_8_content.add_run("本年度坚持以习近平新时代中国特色社会主义思想为指导，全面贯彻落实党的二十大精神及上级党组织工作部署，").font.size = Pt(12)
add_yellow_text(p4_8_content, "扎实推进党建与业务深度融合，充分发挥党组织把方向、管大局、保落实的领导作用").font.bold = True
p4_8_content.add_run("，团结带领全体干部职工凝心聚力、攻坚克难，为公司高质量发展提供坚实的保障，推动科研保障、后勤服务、经营生产、安全保密等工作顺利开展。").font.size = Pt(12)

# 二、2026年度工作计划
section2 = doc.add_heading('', level=1)
run = section2.add_run("二、2026年度工作计划")
run.font.size = Pt(16)

p5_intro = doc.add_paragraph()
p5_intro.add_run("下一年度，董事会将").font.size = Pt(12)
add_yellow_text(p5_intro, "切实履行董事职责，充分发挥决策监管、战略引领作用").font.bold = True
p5_intro.add_run("，立足公司科研保障、物业及总务服务、经营生产等职能定位，坚持党建引领与规范治理并重，").font.size = Pt(12)
add_yellow_text(p5_intro, "更好服务我所发展").font.bold = True
p5_intro.add_run("，结合公司具体业务，").font.size = Pt(12)
add_yellow_text(p5_intro, "全力推进").font.bold = True
p5_intro.add_run("以下重点工作：").font.size = Pt(12)

# 计划详情
plans = [
    ("1. 强化公司治理，规范运营管理行为", [
        "一是严格按照公司章程和上级单位监管要求，履行董事职责，正式行使职权，",
        add_yellow_text(doc.add_paragraph(), "切实维护股东和公司利益"),
        "二是加强与公司领导班子沟通协调，促进决策科学、执行高效、监督到位；",
        "三是推动内部制度完善，关注财务风险、采购及合同风险，",
        add_yellow_text(doc.add_paragraph(), "确保内控评价常态化，公司合规运营")
    ]),
    ("2. 聚焦主责主业，提升服务保障质效", [
        "一是开展全园区满意度测评，覆盖服务保障关键指标，",
        add_yellow_text(doc.add_paragraph(), "建立问题台账明确整改时限"),
        "，持续提升公司服务保障能力；二是加强园区安全守卫监督管理，加强人员专业技能培训，提升各园区应急能力；三是",
        add_yellow_text(doc.add_paragraph(), "推动印刷业务转型升级"),
        "，优化产品结构，实施精益管理，",
        add_yellow_text(doc.add_paragraph(), "聚焦降本增效、提质提效目标"),
        "，推动生产经营精细化升级。"
    ]),
    ("3. 关注队伍建设，凝心聚力服务保障", [
        "一是关心公司人才队伍建设，支持开展特种作业技能、安全保卫、物业管理和精益管理等培训，打造",
        add_yellow_text(doc.add_paragraph(), "专业化经营和服务团队"),
        "；二是关注员工工作生活保障，维护员工合法权益，营造",
        add_yellow_text(doc.add_paragraph(), "爱岗敬业、服务至上、务实担当"),
        "的文化氛围，提升凝聚力和战斗力。"
    ]),
    ('4. 坚持党建引领，把牢正确发展方向', [
        "坚持贯彻落实上级党组织决策部署和工作要求，支持公司党组织发挥",
        add_yellow_text(doc.add_paragraph(), '领导核心作用'),
        "，严格落实「第一议题」制度，采取「课题式调研、项目式研讨」等工作方法，将学习成果转化为",
        add_yellow_text(doc.add_paragraph(), '破解难题的解决方案和优化管理的制度'),
        "，确保公司发展服务于我所整体战略规划。"
    ])
]

for title, contents in plans:
    p_plan_title = doc.add_paragraph()
    run = p_plan_title.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    
    for content in contents:
        if isinstance(content, str):
            doc.add_paragraph(content).paragraph_format.first_line_indent = Cm(0.5)
        # else it's already a paragraph

# 结尾
p_ending = doc.add_paragraph()
p_ending.add_run("\n").font.size = Pt(12)

p_ending2 = doc.add_paragraph()
p_ending2.add_run("2025年面对复杂的内部改革与外部环境，董事专题会各项工作有序推进，公司在战略重组、业务整合、合规治理等方面取得了").font.size = Pt(12)
add_yellow_text(p_ending2, "关键进展").font.bold = True
p_ending2.add_run("。").font.size = Pt(12)

p_ending3 = doc.add_paragraph()
p_ending3.add_run("未来，").font.size = Pt(12)
add_yellow_text(p_ending3, "董事及公司高级管理人员将继续恪尽职守、勤勉履职").font.bold = True
p_ending3.add_run('，团结带领全体员工，聚焦公司「十五五」规划蓝图，').font.size = Pt(12)
add_yellow_text(p_ending3, "攻坚克难，锐意进取，全力实现公司发展目标").font.bold = True
p_ending3.add_run("，").font.size = Pt(12)
add_yellow_text(p_ending3, "回报全体股东与公司员工的信任与支持").font.bold = True
p_ending3.add_run("。").font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# 署名
p_sig = doc.add_paragraph()
p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_sig.add_run("上海麒麟士企业发展有限公司董事会").font.size = Pt(12)

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_date.add_run("2026年4月").font.size = Pt(12)

# 保存
output_path = '/Users/mac/Desktop/小燕子成果文件库/文档资料/2025年度董事工作报告（优化版-标注）.docx'
doc.save(output_path)
print(f"优化版报告已生成：{output_path}")
