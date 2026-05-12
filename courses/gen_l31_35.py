#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Aesop's Fables - English Learning L31-35 (Advanced)
Generated for 三先生
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name='Times New Roman', size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if level == 1:
        set_font(run, 'Arial', 18, bold=True, color=(0, 51, 102))
    elif level == 2:
        set_font(run, 'Arial', 14, bold=True, color=(0, 102, 153))
    else:
        set_font(run, 'Arial', 12, bold=True, color=(51, 51, 51))
    return p

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, 'Arial', 12, bold=True, color=(102, 51, 0))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bilingual_paragraph(doc, en_text, zh_text):
    p1 = doc.add_paragraph()
    run1 = p1.add_run(en_text)
    set_font(run1, 'Georgia', 11, color=(0, 0, 0))
    p1.paragraph_format.space_after = Pt(6)
    p1.paragraph_format.line_spacing = 1.5

    p2 = doc.add_paragraph()
    run2 = p2.add_run(zh_text)
    set_font(run2, '宋体', 11, color=(80, 80, 80))
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.5

def add_vocabulary_table(doc, vocab_list):
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    headers = ['Word', 'Phonetic', 'POS', 'Chinese Meaning']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, 'Arial', 10, bold=True, color=(255, 255, 255))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = header_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '003366')
        shd.set(qn('w:color'), '003366')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for word, phonetic, pos, meaning in vocab_list:
        row_cells = table.add_row().cells
        row_cells[0].text = word
        row_cells[1].text = phonetic
        row_cells[2].text = pos
        row_cells[3].text = meaning

        for i, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, 'Arial', 10)
                if i == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif i == 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return table

def add_grammar_section(doc, grammar_points):
    for point in grammar_points:
        p = doc.add_paragraph()
        run = p.add_run(f"• {point}")
        set_font(run, 'Arial', 11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)

def add_exercises(doc, exercises):
    for i, exercise in enumerate(exercises, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {exercise}")
        set_font(run, 'Arial', 11)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(0.5)

def create_document():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # TITLE PAGE
    add_heading(doc, "English Learning Course", level=1)
    add_heading(doc, "Level 31-35: Advanced", level=2)
    add_heading(doc, "Aesop's Fables 伊索寓言", level=2)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("📚 5 Classic Fables with Comprehensive Learning Materials 📚")
    set_font(run, 'Arial', 12, color=(51, 51, 51))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Each story includes:")
    set_font(run, 'Arial', 11, color=(80, 80, 80))

    features = ["✦ Bilingual Text (English & Chinese)", "✦ Vocabulary Glossary", "✦ Grammar Points", "✦ Practice Exercises"]
    for feature in features:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(feature)
        set_font(run, 'Arial', 11, color=(80, 80, 80))

    doc.add_page_break()

    # STORIES
    stories = [
        {
            'title_en': 'The Shepherd Boy and the Wolf',
            'title_zh': '牧童与狼',
            'level': 'Level 31',
            'text': """A shepherd boy tended his sheep near a village. He often cried "Wolf! Wolf!" to make the villagers run to help him chase the wolf away.

When they came, they found no wolf, and the boy laughed at their anger. He did this several times.

One day, a wolf really came. The boy cried "Wolf! Wolf!" as loud as he could. But this time, nobody came to help him.

The villagers thought he was trying to trick them again, so they did not run to help. The wolf scattered the flock, and the boy learned a bitter lesson.

A liar is not believed even when he speaks the truth.""",

            'translation': """一个牧童在村庄附近放羊。他经常大喊"狼来了！狼来了！"让村民们跑来帮他赶走狼。

当他们跑来时，发现没有狼，牧童就嘲笑他们的愤怒。他这样捉弄了好几次。

有一天，狼真的来了。牧童用尽全力大喊"狼来了！狼来了！"但这一次，没有人跑来帮助他。

村民们以为他又在骗他们，所以没有跑去帮忙。狼吃掉了羊群，牧童得到了惨痛的教训。

说谎的人即使说真话也没有人相信。""",

            'vocab': [
                ('shepherd', '/ˈʃepəd/', 'n.', '牧羊人'),
                ('tend', '/tend/', 'v.', '照料，看管'),
                ('flock', '/flɒk/', 'n.', '羊群'),
                ('trick', '/trɪk/', 'v.', '欺骗，捉弄'),
                ('scatter', '/ˈskætə(r)/', 'v.', '驱散，散开'),
                ('liar', '/ˈlaɪə(r)/', 'n.', '说谎者'),
                ('villager', '/ˈvɪlɪdʒə(r)/', 'n.', '村民'),
                ('bitter', '/ˈbɪtə(r)/', 'adj.', '痛苦的，苦的'),
                ('lesson', '/ˈlesn/', 'n.', '教训，课'),
                ('cry', '/kraɪ/', 'v.', '喊，叫'),
            ],

            'grammar': [
                "Simple Past Tense (一般过去时): 'cried', 'came', 'found' — 描述过去发生的事件",
                "Reported Speech (间接引语): 描述牧童的叫喊内容",
                "'Used to' Structure (used to结构): 'He used to cry Wolf!' — 描述过去的习惯",
                "Conditional Sentences (条件句): 'If they came...' — 描述假设情况",
                "'Even when' Clause (even when从句): 'even when he speaks the truth' — 表示让步",
            ],

            'exercises': [
                "Why didn't the villagers come to help the boy when the wolf really came?",
                "Translate the moral into Chinese: 'A liar is not believed even when he speaks the truth.'",
                "Find three examples of past tense verbs in the story and change them to present tense.",
                "Write a short paragraph about a time someone didn't believe you. Use at least 3 vocabulary words.",
                "What lesson did the shepherd boy learn? Do you think he deserved this outcome? Why?",
            ]
        },

        {
            'title_en': 'The Farmer and His Sons',
            'title_zh': '农夫与儿子们',
            'level': 'Level 32',
            'text': """An old farmer had several sons who were always quarreling. He tried to persuade them to live in harmony, but they would not listen.

One day, he called his sons together. He gave them a bundle of sticks and told them to break it.

Each son tried to break the bundle, but none of them could do it. Then the farmer untied the bundle and gave each son a single stick. This time, they broke them easily.

"My sons," said the farmer, "if you stand together, no one can hurt you. But if you are divided, you will be weak like these single sticks."

United we stand, divided we fall.""",

            'translation': """一位老农夫有几个儿子，他们总是争吵。老农夫试图劝说他们和睦相处，但他们不听。

有一天，他把儿子们叫到一起，给了他们一捆木棍，让他们把木棍折断。

每个儿子都试着折断那捆木棍，但没有一个人能折断。然后老农夫解开绳子，给每个儿子一根单独的棍子。这一次，他们轻松地就折断了。

"我的儿子们，"农夫说，"如果你们团结在一起，没有人会伤害你们。但如果你们分裂了，你们就会像这些单根的棍子一样脆弱。"

团结则存，分裂则亡。""",

            'vocab': [
                ('farmer', '/ˈfɑːmə(r)/', 'n.', '农夫'),
                ('quarrel', '/ˈkwɒrəl/', 'v.', '争吵，吵架'),
                ('persuade', '/pəˈsweɪd/', 'v.', '说服，劝说'),
                ('harmony', '/ˈhɑːməni/', 'n.', '和谐，和睦'),
                ('bundle', '/ˈbʌndl/', 'n.', '捆，扎'),
                ('sticks', '/stɪks/', 'n.', '棍子，树枝'),
                ('divide', '/dɪˈvaɪd/', 'v.', '分开，分裂'),
                ('united', '/juˈnaɪtɪd/', 'adj.', '团结的，联合的'),
                ('weak', '/wiːk/', 'adj.', '虚弱的，脆弱的'),
                ('easily', '/ˈiːzəli/', 'adv.', '容易地，轻松地'),
            ],

            'grammar': [
                "Conditional 'if' Clause (if条件句): 'If you stand together, no one can hurt you' — 描述真实条件",
                "Past Simple vs Past Continuous (一般过去时与过去进行时): 'tried' vs 'were quarreling'",
                "'None of' Structure (none of结构): 'none of them could do it' — 表示全部否定",
                "Imperative Mood (祈使句): 'Stand together!' — 表示命令或劝告",
                "Modal Verb 'can/could' (can/could): 表示能力或可能性",
            ],

            'exercises': [
                "Why did the farmer give his sons a bundle of sticks first, then single sticks?",
                "Translate the moral into Chinese: 'United we stand, divided we fall.'",
                "Find an example of 'none of' in the story and use it in a new sentence.",
                "Write a short essay about why unity is important in a team or family.",
                "What would happen if the brothers continued to quarrel? Write your prediction in English.",
            ]
        },

        {
            'title_en': 'The Wolf in Sheep\'s Clothing',
            'title_zh': '披着羊皮的狼',
            'level': 'Level 33',
            'text': """A wolf decided to disguise himself so that he could steal sheep without being noticed. He found a sheepskin and wore it over his wool.

In this disguise, the wolf walked among the sheep. The shepherd saw him but thought it was a sheep and did not worry. The wolf had an easy time picking out the fattest sheep for his dinner.

This went on for some time. But one day, the shepherd decided to have lamb for dinner and caught the wolf in the act.

The shepherd tied the wolf to a tree and gave him a good beating. The shepherd said: "You wicked creature! You have killed many of my sheep. I should have known you were a wolf when I saw you walking with my sheep."

A villain may disguise himself for a while, but sooner or later, his true nature will be revealed.""",

            'translation': """一只狼决定伪装自己，这样他就能偷羊而不被发现。他找到一张羊皮，披在自己的毛上。

在这个伪装下，狼混进了羊群。牧羊人看到了他，但以为他是一只羊，所以没有担心。狼轻而易举地挑选了最肥的羊作为晚餐。

这样的情况持续了一段时间。但有一天，牧羊人想吃羊羔肉，却撞见了狼正在作案。

牧羊人把狼绑在树上，好好地揍了他一顿。牧羊人说："你这恶棍！你已经咬死了我许多只羊。当我看到你和我的羊走在一起时，我就应该认出你是只狼。"

坏人可能会伪装一段时间，但迟早会露出真面目。""",

            'vocab': [
                ('disguise', '/dɪsˈɡaɪz/', 'v./n.', '伪装，假扮'),
                ('steal', '/stiːl/', 'v.', '偷窃，偷盗'),
                ('sheepskin', '/ˈʃiːpskɪn/', 'n.', '羊皮'),
                ('villain', '/ˈvɪlən/', 'n.', '坏人，恶棍'),
                ('wicked', '/ˈwɪkɪd/', 'adj.', '邪恶的，坏的'),
                ('beating', '/ˈbiːtɪŋ/', 'n.', '殴打，打击'),
                ('reveal', '/rɪˈviːl/', 'v.', '揭示，揭露'),
                ('disguise', '/dɪsˈɡaɪz/', 'v.', '伪装'),
                ('fatter', '/ˈfætə(r)/', 'adj.', '更肥的（fat的比较级）'),
                ('sooner or later', '/ˈsuːnə(r) ɔː(r) ˈleɪtə(r)/', 'phrase', '迟早，早晚'),
            ],

            'grammar': [
                "Past Continuous (过去进行时): 'was walking' — 描述正在进行的动作",
                "Modal Verbs 'can/could' (情态动词): 'could steal' — 表示能力",
                "Reported Speech (间接引语): 描述牧羊人说的话",
                "'Sooner or later' (迟早): 表示某事必然发生",
                "Causative 'have' (使役动词have): 'have lamb for dinner' — 表示让人做某事",
            ],

            'exercises': [
                "Why did the wolf disguise himself? Do you think this was a good plan? Why or why not?",
                "Find and write down the moral of this story in Chinese.",
                "What does 'a wolf in sheep's clothing' mean in English? Can you use it in a sentence?",
                "Identify all the past tense verbs in the second paragraph.",
                "Write a short story about a time someone was not what they seemed to be.",
            ]
        },

        {
            'title_en': 'The Cat and the Birds',
            'title_zh': '猫与鸟',
            'level': 'Level 34',
            'text': """A cat heard that some birds in an aviary were ill. He decided to visit the aviary and pretending to be a doctor, he stood outside the door.

"What a pity! Your illness is very serious," said the cat. "You must take my advice. Come out of those cages and let me prepare some medicine for you."

The birds looked at the cat suspiciously and said nothing.

Then the cat said, "Why don't you come out and walk about? Fresh air will do you good. You have been locked up too long."

But the birds knew that he was not a friend, so they stayed inside and kept silent. The cat soon got tired of waiting and walked away, saying: "All my patients are fools!"

You can fool some people some of the time, but you cannot fool all the people all the time.""",

            'translation': """一只猫听说动物园里的几只鸟生病了。他决定去拜访动物园，假装成医生，站在门外。

"真可惜！你们的病很严重，"猫说。"你们必须听我的建议。飞出笼子，让我给你们准备一些药。"

鸟儿们怀疑地看着猫，什么也没说。

然后猫说："你们为什么不出来走走？新鲜空气对你们有好处。你们被关得太久了。"

但鸟儿们知道猫不是朋友，所以他们待在笼子里保持沉默。猫等了一会儿就厌倦了，走开了，说："我的病人都是傻瓜！"

你可以一时欺骗所有人，也可以永远欺骗一些人，但你不可能永远欺骗所有的人。""",

            'vocab': [
                ('aviary', '/ˈeɪviəri/', 'n.', '动物园，大鸟笼'),
                ('doctor', '/ˈdɒktə(r)/', 'n.', '医生'),
                ('suspiciously', '/səˈspɪʃəsli/', 'adv.', '怀疑地，疑心地'),
                ('cage', '/keɪdʒ/', 'n.', '笼子'),
                ('medicine', '/ˈmedsn/', 'n.', '药，医学'),
                ('patient', '/ˈpeɪʃnt/', 'n.', '病人'),
                ('fool', '/fuːl/', 'n./v.', '傻瓜，愚弄'),
                ('fresh', '/freʃ/', 'adj.', '新鲜的'),
                ('lock up', '/lɒk ʌp/', 'v.', '锁起来，关押'),
                ('get tired of', '/ɡet ˈtaɪəd əv/', 'phrase', '厌倦做某事'),
            ],

            'grammar': [
                "Reported Speech (间接引语): The cat's advice to the birds",
                "Modal Verbs 'must/should' (must/should): 表示义务和建议",
                "'Why don't you' (提建议): 'Why don't you come out?' — 表示建议",
                "Proverb Structure (谚语结构): 'You can fool...but you cannot...' — 表示对比",
                "Past Tense Narration (过去时叙述): 讲述故事的经过",
            ],

            'exercises': [
                "Why did the cat pretend to be a doctor? What was his real intention?",
                "Translate the proverb into Chinese and explain its meaning.",
                "Find two examples of the cat trying to persuade the birds to come out.",
                "Write a short dialogue between two birds discussing whether to trust the cat.",
                "Do you think the cat's plan was clever or foolish? Give reasons for your answer.",
            ]
        },

        {
            'title_en': 'The Two Pots',
            'title_zh': '两只水罐',
            'level': 'Level 35',
            'text': """Two pots, one made of brass and one of clay, were floating down a river together. The clay pot tried to keep away from the brass pot, saying: "Keep your distance, please!"

"Why?" asked the brass pot. "Is there anything wrong with my company?"

"Nothing wrong with you," replied the clay pot. "But if I knock against you, I will break. You are so hard, and I am so fragile. Even if you don't mean to hit me, we might collide, and I would be destroyed."

The brass pot understood and agreed to keep his distance. They floated on in peace.

Even people who mean no harm can cause damage if they are not compatible. The strong should protect the weak, and the weak should avoid those who might hurt them even by accident.""",

            'translation': """两只水罐，一只铜制，一只陶制，一起在河里漂着。陶罐试图远离铜罐，说："请保持距离！"

"为什么？"铜罐问。"我的陪伴有什么问题吗？"

"你没什么问题，"陶罐回答说。"但如果我撞到你，我就会碎。你太硬了，而我太脆弱了。即使你不是故意撞我，我们也可能会相撞，那我就会粉身碎骨。"

铜罐明白了，同意保持距离。它们和平地继续漂流。

即使没有恶意的人，如果彼此不相容，也可能造成伤害。强者应该保护弱者，弱者应该避开那些可能无意间伤害他们的人。""",

            'vocab': [
                ('pot', '/pɒt/', 'n.', '罐子，壶'),
                ('brass', '/brɑːs/', 'n.', '黄铜，铜'),
                ('clay', '/kleɪ/', 'n.', '黏土，陶土'),
                ('float', '/fləʊt/', 'v.', '漂浮，漂流'),
                ('fragile', '/ˈfrædʒaɪl/', 'adj.', '脆弱的，易碎的'),
                ('collide', '/kəˈlaɪd/', 'v.', '相撞，碰撞'),
                ('compatible', '/kəmˈpætəbl/', 'adj.', '兼容的，相容的'),
                ('destroy', '/dɪˈstrɔɪ/', 'v.', '毁坏，破坏'),
                ('keep distance', '/kiːp ˈdɪstəns/', 'phrase', '保持距离'),
                ('by accident', '/baɪ ˈæksɪdənt/', 'phrase', '偶然地，意外地'),
            ],

            'grammar': [
                "Future Tense 'will' (将来时): 'I will break' — 表示自然结果",
                "Modal Verb 'would' (would): 'I would be destroyed' — 表示假设结果",
                "'Even if' Clause (even if从句): 'Even if you don't mean to hit me' — 表示让步",
                "Conditional Sentences (条件句): 'If we collide, I would be destroyed' — 描述假设",
                "Passive Voice (被动语态): 'I would be destroyed' — 表示被动承受",
            ],

            'exercises': [
                "Why did the clay pot want to keep away from the brass pot? Was this a wise decision?",
                "Translate the last sentence of the story into Chinese.",
                "Find an example of 'even if' in the story and write a new sentence using the same structure.",
                "Do you think the clay pot was being unfair to the brass pot? Why or why not?",
                "Write about a situation where two very different people or things tried to work together. What happened?",
            ]
        }
    ]

    for story in stories:
        add_heading(doc, f"{story['level']}", level=2)
        add_heading(doc, f"{story['title_en']}", level=1)
        add_heading(doc, f"《{story['title_zh']}》", level=2)
        doc.add_paragraph()

        add_section_title(doc, "📖 Story Text 故事原文")
        paragraphs = story['text'].strip().split('\n\n')
        zh_paragraphs = story['translation'].strip().split('\n\n')
        for i, (en_p, zh_p) in enumerate(zip(paragraphs, zh_paragraphs)):
            add_bilingual_paragraph(doc, en_p.strip(), zh_p.strip())

        doc.add_paragraph()
        add_section_title(doc, "📚 Vocabulary 词汇表")
        add_vocabulary_table(doc, story['vocab'])
        doc.add_paragraph()
        add_section_title(doc, "📝 Grammar Points 语法点")
        add_grammar_section(doc, story['grammar'])
        doc.add_paragraph()
        add_section_title(doc, "✏️ Practice Exercises 练习题")
        add_exercises(doc, story['exercises'])
        doc.add_page_break()

    # END PAGE
    add_heading(doc, "Congratulations! 恭喜完成 L31-35!", level=1)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🎉 You've completed all 5 Aesop's Fables! 🎉")
    set_font(run, 'Arial', 14, color=(0, 102, 51))

    doc.add_paragraph()
    summary = ["📖 Stories Mastered: 5", "📚 New Words Learned: 50", "✏️ Practice Exercises: 25", "🌟 Keep practicing and enjoy your English learning journey!"]
    for item in summary:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_font(run, 'Arial', 12, color=(51, 51, 51))

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("英语学习 L31-35 | 伊索寓言进阶 | 2026")
    set_font(run, 'Arial', 10, color=(128, 128, 128))

    output_path = '/Users/mac/.openclaw/workspace/courses/英语学习_L31-35_AesopFables.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == '__main__':
    create_document()
    print("Done")
