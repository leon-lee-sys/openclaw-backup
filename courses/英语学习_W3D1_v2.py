#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英语学习课件 - Week 3: Presentation Skills
Day 1: How to Start a Presentation
中英文对照版
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = '/Users/mac/.openclaw/workspace/courses/英语学习_W3D1_PresentationStart.docx'

doc = Document()

# ===== 标题 =====
title = doc.add_heading('📚 Week 3: Presentation Skills', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_heading('Day 1: 如何开场演讲（中英文对照）', level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ===== 今日听力材料 =====
doc.add_heading('🎧 今日听力原文｜Today\'s Listening', level=2)

p = doc.add_paragraph()
run = p.add_run('Listen carefully and practice reading aloud.')
run.italic = True

doc.add_paragraph()

# 英文原文
quote_en = doc.add_paragraph()
quote_en.paragraph_format.left_indent = Inches(0.3)
run = quote_en.add_run('"Good morning, everyone. Thank you for being here today. I\'m delighted to share with you our company\'s vision for the next decade. Over the next 20 minutes, I\'ll walk you through our strategic plan, highlight our key achievements, and outline the exciting opportunities ahead. Let\'s begin!"')
run.italic = True
run.font.color.rgb = RGBColor(30, 60, 114)

# 中文翻译
quote_cn = doc.add_paragraph()
quote_cn.paragraph_format.left_indent = Inches(0.3)
run = quote_cn.add_run('"各位早上好。感谢大家今天的到来。我很高兴在这里与大家分享我们公司未来十年的愿景。在接下来的20分钟里，我将向大家介绍我们的战略规划，突出我们的主要成就，并概述未来的机遇。让我们开始吧！"')
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ===== 核心词汇 =====
doc.add_heading('📝 核心词汇｜Key Vocabulary', level=2)

words = [
    ('delighted', '/dɪˈlaɪtɪd/', '高兴的、愉快的', 'I\'m delighted to share with you our vision.', '我很高兴与大家分享我们的愿景。'),
    ('vision', '/ˈvɪʒn/', '愿景、展望', 'Our company has a clear vision for the future.', '我们公司对未来有清晰的愿景。'),
    ('decade', '/ˈdekeɪd/', '十年', 'In the next decade, we aim to double revenue.', '在接下来的十年里，我们的目标是收入翻番。'),
    ('outline', '/ˈaʊtlɪn/', '概述、概括', 'Let me outline our key strategies.', '让我概述我们的关键策略。'),
    ('highlight', '/ˈhaɪlaɪt/', '突出、强调', 'Let me highlight our major achievements.', '让我突出我们的主要成就。'),
    ('strategic', '/strəˈtiːdʒɪk/', '战略性的', 'This is a strategic decision for our company.', '这是我们公司的一项战略性决策。'),
    ('opportunity', '/ˌɒpəˈtjuːnɪtɪ/', '机会', 'There are many opportunities in this market.', '这个市场有很多机会。'),
    ('presentation', '/ˌprezənˈteɪʃn/', '演讲、演示', 'This presentation will last 30 minutes.', '这个演讲将持续30分钟。'),
]

for word, phonetic, meaning, example_en, example_cn in words:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'{phonetic} — {meaning}')
    run.font.size = Pt(11)

    p_en = doc.add_paragraph()
    p_en.paragraph_format.left_indent = Inches(0.3)
    run_en = p_en.add_run(f'  例句：{example_en}')
    run_en.italic = True
    run_en.font.size = Pt(10)
    run_cn = p_en.add_run(f'  中文：{example_cn}')
    run_cn.font.size = Pt(10)
    run_cn.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ===== 实用表达 =====
doc.add_heading('🗣️ 今日实用表达｜Useful Expressions', level=2)

# 开场白
doc.add_heading('1. 开场白｜Greetings', level=3)
greetings = [
    ('Good morning / Good afternoon / Good evening, everyone.', '各位早上好 / 下午好 / 晚上好。'),
    ('Thank you for being here today.', '感谢大家今天的到来。'),
    ('Thank you for joining us.', '感谢大家参加。'),
    ('I\'m delighted / honored to be here.', '我很高兴 / 荣幸在这里。'),
    ('It\'s great to see so many of you.', '很高兴见到这么多朋友。'),
]
for en, cn in greetings:
    p = doc.add_paragraph()
    run = p.add_run(f'• {en}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# 介绍时长
doc.add_heading('2. 介绍时长｜Setting the Agenda', level=3)
agenda = [
    ('Over the next [X] minutes, I\'ll share...', '在接下来的[X]分钟里，我将分享……'),
    ('This presentation will take about [X] minutes.', '这个演讲大约需要[X]分钟。'),
    ('I\'ll keep this brief — just 15 minutes.', '我会简要说明——只需15分钟。'),
    ('We have about [X] minutes, so let\'s get started.', '我们大约有[X]分钟，我们开始吧。'),
]
for en, cn in agenda:
    p = doc.add_paragraph()
    run = p.add_run(f'• {en}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# 预告内容
doc.add_heading('3. 预告内容｜Preview', level=3)
preview = [
    ('I\'ll walk you through our plans for...', '我将向大家介绍我们的……计划。'),
    ('Let me begin by giving you an overview of...', '让我首先给大家概述一下……'),
    ('First, let me tell you about...', '首先，让我告诉大家……'),
    ('I\'d like to start with...', '我想从……开始。'),
    ('Our agenda today includes three parts: ...', '今天的议程包括三部分：……'),
]
for en, cn in preview:
    p = doc.add_paragraph()
    run = p.add_run(f'• {en}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# 结束开场
doc.add_heading('4. 结束开场白｜Getting Started', level=3)
starters = [
    ('Let\'s begin!', '让我们开始吧！'),
    ('Shall we get started?', '我们可以开始了吗？'),
    ('Without further ado, let\'s dive in.', '闲话少说，让我们开始吧。'),
    ('Let\'s get straight to it.', '让我们直接进入主题。'),
    ('So, let\'s get started.', '好，让我们开始。'),
]
for en, cn in starters:
    p = doc.add_paragraph()
    run = p.add_run(f'• {en}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== 今日练习 =====
doc.add_heading('✍️ 今日口语练习｜Speaking Practice', level=2)

p = doc.add_paragraph()
run = p.add_run('任务：')
run.bold = True
run = p.add_run('用英语做一个2分钟的开场白，介绍你正在做的一个项目/计划。')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('参考句型：')
run.bold = True
run.underline = True

sentences = [
    ('Good morning, everyone. Thank you for joining me today.', '各位早上好。感谢大家今天的到来。'),
    ('I\'m delighted to share with you our project on...', '我很高兴与大家分享我们的……项目。'),
    ('Over the next 5 minutes, I\'ll walk you through...', '在接下来的5分钟里，我将向大家介绍……'),
    ('First, let me give you a quick overview of what we\'re doing.', '首先，让我给大家快速概述一下我们在做什么。'),
    ('Then I\'ll highlight our key milestones.', '然后，我将突出我们的关键里程碑。'),
    ('Finally, I\'ll outline the next steps.', '最后，我将概述下一步的计划。'),
    ('Let\'s begin!', '让我们开始吧！'),
]
for en, cn in sentences:
    p = doc.add_paragraph()
    run = p.add_run(f'• {en}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== 小提示 =====
doc.add_heading('💡 小提示｜Tips', level=2)

p = doc.add_paragraph()
run = p.add_run('演讲开场三感：')
run.bold = True
run.underline = True

tips = [
    ('存在感', '感谢听众到场，让他们感到被重视', 'Make people feel valued by thanking them for coming.'),
    ('价值感', '让听众知道今天能收获什么', 'Tell the audience what they will gain from listening.'),
    ('节奏感', '预告时长和内容结构，建立期待', 'Set expectations by announcing the duration and structure.'),
]

for title_tip, desc_cn, desc_en in tips:
    p = doc.add_paragraph()
    run = p.add_run(f'🎯 {title_tip}：')
    run.bold = True
    run = p.add_run(f'{desc_cn}')
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ===== 语法加油站 =====
doc.add_heading('🔧 语法加油站｜Grammar Focus', level=2)

p = doc.add_paragraph()
run = p.add_run('表达"将要"的三种方式：')
run.bold = True

grammar = [
    ('will + 动原', '一般性预测或临时决定', 'I\'ll share the results with you.', '我将和大家分享结果。'),
    ('be going to + 动原', '计划好的事或有依据的预测', 'I\'m going to present three key points.', '我将介绍三个关键点。'),
    ('be about to + 动原', '即将发生（very soon）', 'We\'re about to start the presentation.', '我们即将开始演讲。'),
]

for name, use, example_en, example_cn in grammar:
    p = doc.add_paragraph()
    run = p.add_run(f'• {name}')
    run.bold = True
    run = p.add_run(f' — {use}')
    p_en = doc.add_paragraph()
    p_en.paragraph_format.left_indent = Inches(0.3)
    run_en = p_en.add_run(f'  例句：{example_en}')
    run_en.italic = True
    run_en.font.size = Pt(10)
    run_cn = p_en.add_run(f'  中文：{example_cn}')
    run_cn.font.size = Pt(10)
    run_cn.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('🌟 Week 3 - Day 1 | 每天半小时，坚持就是胜利！')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUT)
print(f"Saved: {OUT}")
