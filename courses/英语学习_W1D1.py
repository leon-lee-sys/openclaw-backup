#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英语学习课件 - Week 1: Beginner Basics
Day 1: Simple Present & Daily Life
初级：从最简单的日常英语开始
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = '/Users/mac/.openclaw/workspace/courses/英语学习_W1D1_Beginner.docx'

doc = Document()

# ===== 标题 =====
title = doc.add_heading('📚 Week 1: Beginner Basics', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_heading('Day 1: Simple Life, Simple English（简单生活，简单英语）', level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

level_note = doc.add_paragraph()
level_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = level_note.add_run('🎯 难度：初级 Beginner | ⏱ 建议学习时间：30分钟')
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ===== 今日主题 =====
doc.add_heading('🌟 今日主题：Daily Routine（日常生活）', level=2)

intro = doc.add_paragraph()
run = intro.add_run('为什么从"日常生活"开始？因为这些句子最简单、最实用！')
run.italic = True

doc.add_paragraph()

# ===== Part 1: 核心句子 =====
doc.add_heading('📖 Part 1: 10句日常英语', level=2)

sentences = [
    ('I wake up at 7 every morning.', '我每天早上7点起床。'),
    ('I have breakfast at home.', '我在家吃早餐。'),
    ('I go to work by car.', '我开车去上班。'),
    ('I start work at 9 o\'clock.', '我9点开始工作。'),
    ('I have lunch at noon.', '我中午吃午饭。'),
    ('I finish work at 6 o\'clock.', '我6点下班。'),
    ('I cook dinner at home.', '我在家做晚饭。'),
    ('I watch TV in the evening.', '我晚上看电视。'),
    ('I go to bed at 10.', '我10点上床睡觉。'),
    ('It is a good day.', '今天是美好的一天。'),
]

for i, (en, cn) in enumerate(sentences, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {en}')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'   中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 2: 核心词汇 =====
doc.add_heading('📝 Part 2: 核心词汇', level=2)

words = [
    ('wake up', '-', '醒来、起床', 'I wake up early every day.', '我每天早起。'),
    ('breakfast', '/ˈbrekfəst/', '早餐', 'I have breakfast at 7:30.', '我7:30吃早餐。'),
    ('lunch', '/lʌntʃ/', '午餐', 'I have lunch together.', '我们一起吃午饭吧。'),
    ('dinner', '/ˈdɪnə/', '晚餐', 'Dinner is ready!', '晚餐准备好了！'),
    ('work', '/wɜːk/', '工作', 'I work at an office.', '我在办公室工作。'),
    ('home', '/həʊm/', '家', 'I go home at 6.', '我6点回家。'),
    ('bed', '/bed/', '床', 'Time for bed.', '该睡觉了。'),
    ('TV', '/ˌtiːˈviː/', '电视', 'I watch TV every night.', '我每晚看电视。'),
]

for word, phonetic, meaning, example_en, example_cn in words:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'{phonetic} — {meaning}')
    run.font.size = Pt(11)

    p_en = doc.add_paragraph()
    p_en.paragraph_format.left_indent = Inches(0.3)
    run_en = p_en.add_run(f'  例句：{example_en}')
    run_en.italic = True
    run_en.font.size = Pt(10)
    run_cn = p_en.add_run(f'  中文：{example_cn}')
    run_cn.font.size = Pt(10)
    run_cn.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ===== Part 3: 语法 =====
doc.add_heading('🔧 Part 3: 语法很简单 — 一般现在时', level=2)

p = doc.add_paragraph()
run = p.add_run('一般现在时 = 主语 + 动词原形（第三人称单数加-s）')
run.bold = True
run.font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

grammar = [
    ('I/You/We/They', 'work', '我/你/我们/他们', '工作'),
    ('He/She/It', 'works', '他/她/它', '工作（第三人称单数）'),
]

for subject, verb, cn_subj, cn_verb in grammar:
    p = doc.add_paragraph()
    run = p.add_run(f'• {subject} {verb}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  中文：{cn_subj} {cn_verb}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 4: 练习 =====
doc.add_heading('✍️ Part 4: 今日口语练习', level=2)

p = doc.add_paragraph()
run = p.add_run('任务：用英语介绍你的一天（至少5句）')
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('参考句型：')
run.bold = True
run.underline = True

practice = [
    ('I wake up at [时间].', '我每天[时间]起床。'),
    ('I have breakfast/lunch/dinner at [地点].', '我在[地点]吃早餐/午餐/晚餐。'),
    ('I go to work by [交通方式].', '我乘[交通工具]去上班。'),
    ('I start work at [时间].', '我[时间]开始工作。'),
    ('I finish work at [时间].', '我[时间]下班。'),
    ('In the evening, I [活动].', '晚上，我[活动]。'),
    ('I go to bed at [时间].', '我[时间]睡觉。'),
]

for en, cn in practice:
    p = doc.add_paragraph()
    run = p.add_run(f'• {en}')
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 5: 小故事 =====
doc.add_heading('📖 Part 5: 简单小故事', level=2)

story_en = doc.add_paragraph()
story_en.paragraph_format.left_indent = Inches(0.3)
run = story_en.add_run('My Day')
run.bold = True
run.font.color.rgb = RGBColor(30, 60, 114)

story_en2 = doc.add_paragraph()
story_en2.paragraph_format.left_indent = Inches(0.3)
run2 = story_en2.add_run('I am Tom. I am 30 years old. I work in a big city. Every day, I wake up at 6:30. I have coffee and toast for breakfast. Then I drive to work. I love my job. In the evening, I cook dinner. After dinner, I read a book. At 10 o\'clock, I go to bed. Life is simple, but I am happy.')
run2.font.size = Pt(11)

doc.add_paragraph()

story_cn = doc.add_paragraph()
story_cn.paragraph_format.left_indent = Inches(0.3)
run = story_cn.add_run('中文翻译：')
run.bold = True
run.font.color.rgb = RGBColor(80, 80, 80)

story_cn2 = doc.add_paragraph()
story_cn2.paragraph_format.left_indent = Inches(0.3)
run2 = story_cn2.add_run('我叫Tom，30岁。我在大城市工作。每天早上6:30起床。早餐喝咖啡吃吐司。然后我开车去上班。我热爱我的工作。晚上我做晚饭。晚饭后我看书。10点上床睡觉。生活很简单，但我很快乐。')
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ===== 今日总结 =====
doc.add_heading('📋 今日总结', level=2)

summary = [
    ('词汇', '8个日常生活词汇'),
    ('句型', '10句一般现在时日常表达'),
    ('语法', '一般现在时的基本用法'),
    ('练习', '用英语介绍自己的一天'),
]

for item, desc in summary:
    p = doc.add_paragraph()
    run = p.add_run(f'✅ {item}：')
    run.bold = True
    run = p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('🌟 Week 1 - Day 1 | 初级难度 | 每天半小时，坚持就是胜利！')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUT)
print(f"Saved: {OUT}")
