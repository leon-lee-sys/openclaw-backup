#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英语学习课件 - Week 3: Presentation Skills (演讲汇报)
Day 1: How to Start a Presentation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUT = '/Users/mac/.openclaw/workspace/courses/英语学习_W3D1_PresentationStart.docx'

doc = Document()

# 标题
title = doc.add_heading('📚 Week 3: Presentation Skills', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_heading('Day 1: How to Start a Presentation（如何开场演讲）', level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ===== 今日听力材料 =====
h = doc.add_heading('🎧 今日听力材料', level=2)

p = doc.add_paragraph()
run = p.add_run('Listen carefully and practice reading aloud:')
run.italic = True

doc.add_paragraph()

# 引用框
quote = doc.add_paragraph()
quote.paragraph_format.left_indent = Inches(0.5)
run = quote.add_run('"Good morning, everyone. Thank you for being here today. I\'m delighted to share with you our company\'s vision for the next decade. Over the next 20 minutes, I\'ll walk you through our strategic plan, highlight our key achievements, and outline the exciting opportunities ahead. Let\'s begin!"')
run.italic = True
run.font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

# ===== 核心词汇 =====
doc.add_heading('📝 核心词汇', level=2)

words = [
    ('delighted', '/dɪˈlaɪtɪd/', '高兴的、愉快的', 'I\'m delighted to share with you our vision.'),
    ('vision', '/ˈvɪʒn/', '愿景、展望', 'Our company has a clear vision for the future.'),
    ('decade', '/ˈdekeɪd/', '十年', 'In the next decade, we aim to double revenue.'),
    ('outline', '/ˈaʊtlɪn/', '概述、概括', 'Let me outline our key strategies.'),
    ('highlight', '/ˈhaɪlaɪt/', '突出、强调', 'Let me highlight our major achievements.'),
    ('strategic', '/strəˈtiːdʒɪk/', '战略性的', 'This is a strategic decision for our company.'),
    ('opportunity', '/ˌɒpəˈtjuːnɪtɪ/', '机会', 'There are many opportunities in this market.'),
    ('presentation', '/ˌprezənˈteɪʃn/', '演讲、演示', 'This presentation will last 30 minutes.'),
]

for word, phonetic, meaning, example in words:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run2 = p.add_run(f'{phonetic} — {meaning}')
    run2.font.size = Pt(11)
    p2 = doc.add_paragraph()
    run3 = p2.add_run(f'  例句：{example}')
    run3.italic = True
    run3.font.size = Pt(10)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== 实用表达 =====
doc.add_heading('🗣️ 今日实用表达', level=2)

# 开场白
doc.add_heading('1. 开场白（Greetings）', level=3)
greetings = [
    'Good morning / Good afternoon / Good evening, everyone.',
    'Thank you for being here today.',
    'Thank you for joining us.',
    'I\'m delighted / honored to be here.',
    'It\'s great to see so many of you.',
]
for g in greetings:
    doc.add_paragraph(f'• {g}', style='List Bullet')

doc.add_paragraph()

# 介绍时长
doc.add_heading('2. 介绍时长（Setting the Agenda）', level=3)
agenda = [
    'Over the next [X] minutes, I\'ll share...',
    'This presentation will take about [X] minutes.',
    'I\'ll keep this brief — just 15 minutes.',
    'We have about [X] minutes, so let\'s get started.',
]
for a in agenda:
    doc.add_paragraph(f'• {a}', style='List Bullet')

doc.add_paragraph()

# 预告内容
doc.add_heading('3. 预告内容（Preview）', level=3)
preview = [
    'I\'ll walk you through our plans for...',
    'Let me begin by giving you an overview of...',
    'First, let me tell you about...',
    'I\'d like to start with...',
    'Our agenda today includes three parts: ...',
]
for p_item in preview:
    doc.add_paragraph(f'• {p_item}', style='List Bullet')

doc.add_paragraph()

# 结束开场
doc.add_heading('4. 结束开场白（Getting Started）', level=3)
starters = [
    'Let\'s begin!',
    'Shall we get started?',
    'Without further ado, let\'s dive in.',
    'Let\'s get straight to it.',
    'So, let\'s get started.',
]
for s in starters:
    doc.add_paragraph(f'• {s}', style='List Bullet')

doc.add_paragraph()

# ===== 今日练习 =====
doc.add_heading('✍️ 今日口语练习', level=2)

p = doc.add_paragraph()
run = p.add_run('任务：')
run.bold = True
run = p.add_run('用英语做一个2分钟的开场白，介绍你正在做的一个项目/计划。')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('参考句型：')
run.bold = True
run.underline = True

sentences = [
    'Good morning, everyone. Thank you for joining me today.',
    'I\'m delighted to share with you our project on...',
    'Over the next 5 minutes, I\'ll walk you through...',
    'First, let me give you a quick overview of what we\'re doing.',
    'Then I\'ll highlight our key milestones.',
    'Finally, I\'ll outline the next steps.',
    'Let\'s begin!',
]
for s in sentences:
    doc.add_paragraph(f'• {s}', style='List Bullet')

doc.add_paragraph()

# ===== 小提示 =====
doc.add_heading('💡 小提示：演讲开场三感', level=2)

tips = [
    ('存在感', '感谢听众到场，让他们感到被重视'),
    ('价值感', '让听众知道今天能收获什么'),
    ('节奏感', '预告时长和内容结构，建立期待'),
]

for title_tip, desc in tips:
    p = doc.add_paragraph()
    run = p.add_run(f'🎯 {title_tip}：')
    run.bold = True
    run = p.add_run(desc)
    run.font.size = Pt(11)

doc.add_paragraph()

# ===== 语法加油站 =====
doc.add_heading('🔧 语法加油站', level=2)

p = doc.add_paragraph()
run = p.add_run('表达"将要"的三种方式：')

grammar = [
    ('will + 动原', '一般性预测或临时决定', 'I\'ll share the results with you.'),
    ('be going to + 动原', '计划好的事或有依据的预测', 'I\'m going to present three key points.'),
    ('be about to + 动原', '即将发生（very soon）', 'We\'re about to start the presentation.'),
]

for name, use, example in grammar:
    p = doc.add_paragraph()
    run = p.add_run(f'• {name}：')
    run.bold = True
    run = p.add_run(f'{use} → {example}')

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('🌟 Week 3 - Day 1 | 每天半小时，坚持就是胜利！')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUT)
print(f"Saved: {OUT}")
