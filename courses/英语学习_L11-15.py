#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英语学习材料 · 中级 Lessons 11-15
伊索寓言 · Aesop's Fables
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = '/Users/mac/.openclaw/workspace/courses/英语学习_L11-15_AesopFables.docx'

doc = Document()

# ===== 标题 =====
title = doc.add_heading('📚 英语学习材料 · 中级 Lessons 11-15', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('伊索寓言 Aesop\'s Fables | 每天半小时，坚持就是胜利 🐦')
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ===== Lesson 11 =====
doc.add_heading('📖 Lesson 11: The Bat, the Birds, and the Beasts', level=2)
p = doc.add_paragraph()
run = p.add_run('蝙蝠、鸟类与兽类')
run.italic = True

doc.add_paragraph()

story = """A great war was about to begin between the birds and the beasts. The bat did not know which side to join.

When the birds were winning, he would fly with them. When the beasts were winning, he would run with them.

But finally, the war ended. The birds and the beasts made peace. They both turned against the bat.

"You tricked us both," they said. "You are neither a bird nor a beast. You are a traitor."

The bat was too ashamed to show his face. He hid in a dark cave forever.

【中文翻译】
一场大战在鸟类和兽类之间爆发了。蝙蝠不知道该加入哪一边。

每当鸟类获胜，他就飞过去帮他们；每当兽类获胜，他就跑过去帮他们。

但最终，战争结束了。鸟类和兽类讲和了，但他们都转身对付蝙蝠。

"你两边都欺骗了，"他们说。"你既不是鸟也不是兽。你是个叛徒。"

蝙蝠羞得无脸见人，躲进黑暗的洞穴里，再也不出来了。"""

for line in story.split('\n'):
    if line.strip():
        p = doc.add_paragraph(line.strip())
        if '【中文翻译】' in line:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

vocab = [
    ('bat', '/bæt/', 'n. 蝙蝠'),
    ('beast', '/biːst/', 'n. 野兽'),
    ('war', '/wɔːr/', 'n. 战争'),
    ('trick', '/trɪk/', 'v. 欺骗'),
    ('neither...nor', '/ˈnaɪðər/', '既不...也不'),
    ('traitor', '/ˈtreɪtər/', 'n. 叛徒'),
    ('ashamed', '/əˈʃeɪmd/', 'adj. 羞耻的'),
    ('cave', '/keɪv/', 'n. 洞穴'),
]

p = doc.add_paragraph()
run = p.add_run('核心词汇：')
run.bold = True

for word, phonetic, meaning in vocab:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'{phonetic} — {meaning}')
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

grammar = """语法要点：
1. neither...nor...（既不...也不...）
   例：He is neither a bird nor a beast.

2. turn against（转向对付）
   例：They both turned against the bat."""

p = doc.add_paragraph()
run = p.add_run(grammar)
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ===== Lesson 12 =====
doc.add_heading('📖 Lesson 12: The Wolf in Sheep\'s Clothing', level=2)
p = doc.add_paragraph()
run = p.add_run('披着羊皮的狼')
run.italic = True

doc.add_paragraph()

story = """A wolf wanted to eat the sheep. But he was afraid of the sheepdog.

One day, the wolf found a sheepskin. He put it on and mixed with the sheep.

The shepherd came out. He did not see the wolf. He only saw the sheep.

The wolf enjoyed eating the sheep every night. No one suspected him.

But one night, the wolf laughed too loudly. The shepherd recognized his voice. He caught the wolf and punished him.

【中文翻译】
一只狼想吃羊，但他害怕牧羊犬。

有一天，狼发现了一张羊皮。他把它披在身上，混进了羊群。

牧羊人出来了。他没看见狼，只看见了羊。

狼每天晚上都偷吃羊。没有人怀疑他。

但有一天晚上，狼笑得太响了。牧羊人认出了他的声音，抓住了狼并惩罚了他。"""

for line in story.split('\n'):
    if line.strip():
        p = doc.add_paragraph(line.strip())
        if '【中文翻译】' in line:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

vocab = [
    ('sheepskin', '/ˈʃiːpskɪn/', 'n. 羊皮'),
    ('shepherd', '/ˈʃepərd/', 'n. 牧羊人'),
    ('suspect', '/səˈspekt/', 'v. 怀疑'),
    ('recognize', '/ˈrekəɡnaɪz/', 'v. 认出'),
    ('punish', '/ˈpʌnɪʃ/', 'v. 惩罚'),
    ('mix', '/mɪks/', 'v. 混合'),
    ('laugh', '/læf/', 'v. 笑'),
    ('voice', '/vɔɪs/', 'n. 声音'),
]

p = doc.add_paragraph()
run = p.add_run('核心词汇：')
run.bold = True

for word, phonetic, meaning in vocab:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'{phonetic} — {meaning}')
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

grammar = """语法要点：
1. want to + 动原（想要...）
   例：A wolf wanted to eat the sheep.

2. be afraid of（害怕...）
   例：He was afraid of the sheepdog."""

p = doc.add_paragraph()
run = p.add_run(grammar)
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ===== Lesson 13 =====
doc.add_heading('📖 Lesson 13: The Two Pots', level=2)
p = doc.add_paragraph()
run = p.add_run('两只锅')
run.italic = True

doc.add_paragraph()

story = """Two pots met in a river. One was made of brass, and the other was made of clay.

"Be careful," said the clay pot to the brass pot. "We might hit each other."

"Don't worry," said the brass pot. "I am strong. I can protect myself."

Just then, the river carried them against each other. The clay pot broke into pieces.

The brass pot was sorry for his friend. "Pride is not a good thing," he said.

【中文翻译】
两只锅在河里相遇了。一只是黄铜做的，一只是陶土做的。

"小心点，"陶土锅对黄铜锅说。"我们可能会撞到彼此。"

"别担心，"黄铜锅说。"我很结实，我能保护自己。"

就在这时，河水把他们冲向彼此。陶土锅碎成了碎片。

黄铜锅为他的朋友感到惋惜。"骄傲不是好事，"他说。"""

for line in story.split('\n'):
    if line.strip():
        p = doc.add_paragraph(line.strip())
        if '【中文翻译】' in line:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

vocab = [
    ('pot', '/pɒt/', 'n. 锅'),
    ('brass', '/bræs/', 'n. 黄铜'),
    ('clay', '/kleɪ/', 'n. 陶土'),
    ('careful', '/ˈkeərfəl/', 'adj. 小心的'),
    ('protect', '/prəˈtekt/', 'v. 保护'),
    ('pride', '/praɪd/', 'n. 骄傲'),
    ('break (broke, broken)', '/breɪk/', 'v. 打破'),
    ('river', '/ˈrɪvər/', 'n. 河流'),
]

p = doc.add_paragraph()
run = p.add_run('核心词汇：')
run.bold = True

for word, phonetic, meaning in vocab:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'{phonetic} — {meaning}')
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

grammar = """语法要点：
1. might + 动原（可能...）
   例：We might hit each other.

2. be sorry for（为...感到惋惜）
   例：The brass pot was sorry for his friend."""

p = doc.add_paragraph()
run = p.add_run(grammar)
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ===== Lesson 14 =====
doc.add_heading('📖 Lesson 14: The Astronomer', level=2)
p = doc.add_paragraph()
run = p.add_run('天文学家')
run.italic = True

doc.add_paragraph()

story = """An astronomer loved to study the stars. Every night, he would go to the roof of his house and look at the sky.

One night, while he was looking at the stars, he fell into a hole. He hurt his leg badly.

A kind man passing by helped him up. "You should watch your step," said the man. "You are looking at the sky but not at the ground."

The astronomer learned a lesson: we should not be so absorbed in one thing that we forget everything else.

【中文翻译】
一位天文学家喜欢研究星星。每天晚上，他都会爬上自家屋顶仰望天空。

一天晚上，当他正在看星星时，掉进了一个坑里。他的腿伤得很重。

一个好心人路过，把他扶了起来。"你应该看着脚下，"那人说。"你在看天，但没看地。"

天文学家得到了一个教训：我们不应该太专注于一件事，以至于忘记了其他一切。"""

for line in story.split('\n'):
    if line.strip():
        p = doc.add_paragraph(line.strip())
        if '【中文翻译】' in line:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

vocab = [
    ('astronomer', '/əˈstrɒnəmər/', 'n. 天文学家'),
    ('star', '/stɑːr/', 'n. 星星'),
    ('roof', '/ruːf/', 'n. 屋顶'),
    ('hole', '/həʊl/', 'n. 洞'),
    ('hurt', '/hɜːrt/', 'v. 伤害'),
    ('absorbed', '/əbˈzɔːrbd/', 'adj. 专注的'),
    ('lesson', '/ˈlesən/', 'n. 教训'),
    ('ground', '/ɡraʊnd/', 'n. 地面'),
]

p = doc.add_paragraph()
run = p.add_run('核心词汇：')
run.bold = True

for word, phonetic, meaning in vocab:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'{phonetic} — {meaning}')
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

grammar = """语法要点：
1. while + 进行时（当...的时候）
   例：While he was looking at the stars, he fell into a hole.

2. so...that...（如此...以至于...）
   例：He was so absorbed that he forgot everything else."""

p = doc.add_paragraph()
run = p.add_run(grammar)
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ===== Lesson 15 =====
doc.add_heading('📖 Lesson 15: The Cat and the Mice', level=2)
p = doc.add_paragraph()
run = p.add_run('猫和老鼠')
run.italic = True

doc.add_paragraph()

story = """There was a cat in a house. The mice were afraid of him. They did not know how to stop him.

The mice called a meeting. "We must do something," they said. "The cat is too clever. We cannot catch him."

One mouse had an idea. "Let us hang a bell around the cat's neck. Then we will hear him coming."

All the mice agreed. But one old mouse asked, "Who will tie the bell to the cat?"

No one answered. The plan was impossible.

【中文翻译】
房子里有一只猫。老鼠们很怕他。他们不知道该怎么阻止他。

老鼠们召开了一次会议。"我们必须做点什么，"他们说。"这只猫太聪明了，我们抓不到他。"

一只老鼠有了个主意。"让我们在猫的脖子上挂一个铃铛。这样我们就能听到他来了。"

所有的老鼠都同意了。但一只年迈的老鼠问道："谁去把铃铛系在猫脖子上？"

没有人回答。这个计划是不可能实现的。"""

for line in story.split('\n'):
    if line.strip():
        p = doc.add_paragraph(line.strip())
        if '【中文翻译】' in line:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

vocab = [
    ('mouse (pl. mice)', '/maʊs/', 'n. 老鼠'),
    ('bell', '/bel/', 'n. 铃铛'),
    ('neck', '/nek/', 'n. 脖子'),
    ('hang', '/hæŋ/', 'v. 悬挂'),
    ('clever', '/ˈklevər/', 'adj. 聪明的'),
    ('meeting', '/ˈmiːtɪŋ/', 'n. 会议'),
    ('impossible', '/ɪmˈpɒsəbl/', 'adj. 不可能的'),
    ('tie', '/taɪ/', 'v. 系、绑'),
]

p = doc.add_paragraph()
run = p.add_run('核心词汇：')
run.bold = True

for word, phonetic, meaning in vocab:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'{phonetic} — {meaning}')
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

grammar = """语法要点：
1. be afraid of（害怕...）
   例：The mice were afraid of him.

2. let us + 动原（让我们...）
   例：Let us hang a bell around the cat's neck."""

p = doc.add_paragraph()
run = p.add_run(grammar)
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ===== 学习指南 =====
doc.add_heading('📝 学习指南', level=2)

guide = [
    ('1. 通读全文', '理解故事内容'),
    ('2. 划出生词', '查词典确认意思'),
    ('3. 重读课文', '翻译成中文'),
    ('4. 背诵核心词汇', '每个单词造一个句子'),
    ('5. 重点语法造句', '用今天学的语法写3个句子'),
    ('6. 朗读课文', '注意发音，可录音自纠'),
]

for step, desc in guide:
    p = doc.add_paragraph()
    run = p.add_run(f'✅ {step}：')
    run.bold = True
    run = p.add_run(desc)

doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('📚 Lessons 11-15 | 伊索寓言中级 | 每天半小时，坚持就是胜利！🐦')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUT)
print(f"Saved: {OUT}")
