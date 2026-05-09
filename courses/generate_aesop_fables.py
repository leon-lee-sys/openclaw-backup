#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Aesop's Fables - English Learning L21-25 (Advanced)
Generated for 三先生
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name='Times New Roman', size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # For Chinese font support
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if level == 1:
        set_font(run, 'Arial', 18, bold=True, color=(0, 51, 102))
    elif level == 2:
        set_font(run, 'Arial', 14, bold=True, color=(0, 102, 153))
    else:
        set_font(run, 'Arial', 12, bold=True, color=(51, 51, 51))
    return p

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, 'Arial', 12, bold=True, color=(102, 51, 0))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_normal_paragraph(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 'Arial', 11)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bilingual_paragraph(doc, en_text, zh_text):
    # English paragraph
    p1 = doc.add_paragraph()
    run1 = p1.add_run(en_text)
    set_font(run1, 'Georgia', 11, color=(0, 0, 0))
    p1.paragraph_format.space_after = Pt(6)
    p1.paragraph_format.line_spacing = 1.5

    # Chinese translation
    p2 = doc.add_paragraph()
    run2 = p2.add_run(zh_text)
    set_font(run2, '宋体', 11, color=(80, 80, 80))
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.5

def add_vocabulary_table(doc, vocab_list):
    """Add vocabulary table with Word, Phonetic, Part of Speech, Chinese Meaning"""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Word', 'Phonetic', 'POS', 'Chinese Meaning']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, 'Arial', 10, bold=True, color=(255, 255, 255))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set header background color
        tc = header_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '003366')
        shd.set(qn('w:color'), '003366')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    # Data rows
    for word, phonetic, pos, meaning in vocab_list:
        row_cells = table.add_row().cells
        row_cells[0].text = word
        row_cells[1].text = phonetic
        row_cells[2].text = pos
        row_cells[3].text = meaning

        for i, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, 'Arial', 10)
                if i == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif i == 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return table

def add_grammar_section(doc, grammar_points):
    for point in grammar_points:
        p = doc.add_paragraph()
        run = p.add_run(f"• {point}")
        set_font(run, 'Arial', 11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)

def add_exercises(doc, exercises):
    for i, exercise in enumerate(exercises, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {exercise}")
        set_font(run, 'Arial', 11)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(0.5)

def create_document():
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ==================== TITLE PAGE ====================
    add_heading(doc, "English Learning Course", level=1)
    add_heading(doc, "Level 21-25: Advanced", level=2)
    add_heading(doc, "Aesop's Fables 伊索寓言", level=2)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("📚 5 Classic Fables with Comprehensive Learning Materials 📚")
    set_font(run, 'Arial', 12, color=(51, 51, 51))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Each story includes:")
    set_font(run, 'Arial', 11, color=(80, 80, 80))

    features = [
        "✦ Bilingual Text (English & Chinese)",
        "✦ Vocabulary Glossary",
        "✦ Grammar Points",
        "✦ Practice Exercises"
    ]
    for feature in features:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(feature)
        set_font(run, 'Arial', 11, color=(80, 80, 80))

    doc.add_page_break()

    # ==================== STORIES ====================
    stories = [
        {
            'title_en': 'The Crow and the Pitcher',
            'title_zh': '乌鸦与水瓶',
            'level': 'Level 21',
            'text': """A crow, half-dead with thirst, came upon a Pitcher. When he tried to drink from it he found that the water was low and he couldn't reach it. He tried, but the neck of the pitcher was too narrow.

The crow thought hard, and then he saw some pebbles lying nearby. He began dropping pebbles into the Pitcher, one by one. With each pebble, the water rose a little higher.

At last, the water rose high enough for the crow to drink. The crow drank the water and flew away satisfied.

A little bit of careful thought and problem-solving can go a long way.""",

            'translation': """一只乌鸦渴得快要死了，来到一个水瓶前。当它试着喝水时，发现水位很低，根本够不着。它尝试了，但瓶口太窄了。

乌鸦苦思冥想，然后看到附近有一些小石子。它开始一颗一颗地把石子扔进水瓶里。每扔一颗石子，水就升高一点。

最后，水位升得足够高了，乌鸦终于喝到了水。它喝完水，心满意足地飞走了。

一点点仔细的思考和问题解决能力可以带来很大的帮助。""",

            'vocab': [
                ('thirst', '/θɜːst/', 'n.', '口渴，渴感'),
                ('pitcher', '/ˈpɪtʃə(r)/', 'n.', '大水罐，浇水壶'),
                ('pebble', '/ˈpebl/', 'n.', '鹅卵石，小石子'),
                ('satisfy', '/ˈsætɪsfaɪ/', 'v.', '使满足，使满意'),
                ('problem-solving', '/ˈprɒbləm ˈsɒlvɪŋ/', 'n.', '解决问题'),
                ('half-dead', '/hɑːf ded/', 'adj.', '奄奄一息的，半死的'),
                ('reach', '/riːtʃ/', 'v.', '够到，触及'),
                ('narrow', '/ˈnærəʊ/', 'adj.', '狭窄的，窄的'),
                ('careful', '/ˈkeəfl/', 'adj.', '仔细的，小心的'),
                ('thought', '/θɔːt/', 'n.', '思考，想法'),
            ],

            'grammar': [
                "Past Tense (过去时): 'came', 'found', 'tried', 'began' — 描述故事中发生的事件",
                "Complex Sentence with 'when' (when引导的复合句): 'When he tried to drink from it...' — 表示时间顺序",
                "Future Tense with 'would' (would表示过去将来时): 'the water would rise' — 描述预料中的结果",
                "'Enough + to V' Structure (enough to 结构): 'high enough for the crow to drink' — 表示程度和结果",
                "Modal Verb 'can' (情态动词): 'can go a long way' — 表示能力或可能性",
            ],

            'exercises': [
                "Translate the following sentence into English: '乌鸦把石子扔进水瓶里，水位慢慢升高了。'",
                "Find and explain the 'when' clause in the story. What does it connect?",
                "What is the moral of this fable? Write your answer in English (at least 2 sentences).",
                "Use 'enough...to...' in a sentence about a different scenario.",
                "Identify all the past tense verbs in the first three sentences.",
            ]
        },

        {
            'title_en': 'The Fox and the Grapes',
            'title_zh': '狐狸与葡萄',
            'level': 'Level 22',
            'text': """One hot summer's day a fox was walking through an orchard when he saw a bunch of beautiful, ripe grapes hanging from a high branch.

"Just the thing to quench my thirst!" said the fox.

He backed up several steps, took a running leap, and jumped. But he missed the grapes by a long way.

He tried again and again, but each time he failed. At last, he gave up. He walked away with his nose in the air, saying: "I'm sure those grapes are sour. It is probably just as well that I did not get them. Who would want to eat sour grapes anyway?"

It is easy to despise what you cannot have.""",

            'translation': """一个炎热的夏日，一只狐狸穿过果园时，看到一串串美丽的、成熟的葡萄挂在高高的枝头。

"这正是解渴的好东西！"狐狸说。

它后退了几步，助跑，然后跳起来。但它离葡萄还差得远呢。

它一次又一次地尝试，但每次都失败了。最后，它放弃了。它仰着头走开了，说："我敢说那些葡萄肯定是酸的。我没有得到它们也许正好。谁会想吃酸的葡萄呢？"

容易鄙视得不到的东西。""",

            'vocab': [
                ('orchard', '/ˈɔːtʃəd/', 'n.', '果园'),
                ('bunch', '/bʌntʃ/', 'n.', '串，束'),
                ('ripe', '/raɪp/', 'adj.', '成熟的'),
                ('quench', '/kwentʃ/', 'v.', '止渴，解渴'),
                ('thirst', '/θɜːst/', 'n.', '口渴'),
                ('leap', '/liːp/', 'n./v.', '跳跃'),
                ('despise', '/dɪˈspaɪz/', 'v.', '鄙视，看不起'),
                ('sour', '/ˈsaʊə(r)/', 'adj.', '酸的'),
                ('probably', '/ˈprɒbəbli/', 'adv.', '大概，也许'),
                ('anyway', '/ˈeniweɪ/', 'adv.', '无论如何，反正'),
            ],

            'grammar': [
                "Direct Speech (直接引语): 'Just the thing to quench my thirst!' — 狐狸说的话",
                "Gerund as Subject (动名词作主语): 'Eating sour grapes is foolish.' — 动名词的用法",
                "Past Continuous (过去进行时): 'was walking' — 描述持续进行的动作",
                "Idiomatic Expression (习惯表达): 'by a long way' — 意为'远未达到'",
                "Reported Speech with 'saying' (saying引导的间接引语): 描述狐狸说的话",
            ],

            'exercises': [
                "What is the Chinese meaning of 'quench my thirst'? Try to use it in an English sentence.",
                "Why did the fox say the grapes were sour? Do you agree with this reasoning?",
                "Find two examples of the past continuous tense in the story.",
                "Translate the moral into Chinese: 'It is easy to despise what you cannot have.'",
                "Write a short paragraph about a time you wanted something but couldn't get it. Use at least 3 vocabulary words from this lesson.",
            ]
        },

        {
            'title_en': 'The Ant and the Grasshopper',
            'title_zh': '蚂蚁与蚱蜢',
            'level': 'Level 23',
            'text': """In a field one summer's day, a grasshopper was hopping about, chirping and singing to its heart's content. An ant passed by, bearing along a huge ear of corn.

"Where are you going with that heavy ear of corn?"

"I am storing it for the winter. I suggest you do the same."

"Why bother about winter?" said the grasshopper. "We have plenty of food at present. Let us sing and dance together."

The ant went on its way and continued its toil. When winter came, the grasshopper was found dead, frozen in the cold snow. But the ant had food stored and was warm in its little home.

There is a time for work and a time for play. Make sure you do not forget to prepare for a rainy day.""",

            'translation': """在一个夏日的田野里，一只蚱蜢正在蹦蹦跳跳，尽情地唱着歌。一只蚂蚁背着巨大的玉米穗从旁边经过。

"你背着这么重的玉米穗要去哪里？"

"我正在储存它准备过冬。我建议你也要这么做。"

"何必操心冬天的事呢？"蚱蜢说。"我们现在有足够的食物。让我们一起唱歌跳舞吧。"

蚂蚁继续走它的路，继续辛勤劳动。冬天来了，蚱蜢被发现冻死在冰冷的雪地里。而蚂蚁储存了食物，暖暖地待在小窝里。

有工作的时间，也有娱乐的时间。千万不要忘记为雨天做准备。""",

            'vocab': [
                ('grasshopper', '/ˈɡræshɒpə(r)/', 'n.', '蚱蜢，蝗虫'),
                ('hop', '/hɒp/', 'v.', '跳跃，跳'),
                ('chirping', '/ˈtʃɜːpɪŋ/', 'n./v.', '唧唧叫，鸣叫'),
                ('ant', '/ænt/', 'n.', '蚂蚁'),
                ('bear', '/beə(r)/', 'v.', '携带，承受'),
                ('ear of corn', '/ɪə əv kɔːn/', 'n.', '玉米穗'),
                ('store', '/stɔː(r)/', 'v.', '储存，储藏'),
                ('toil', '/tɔɪl/', 'n./v.', '辛苦劳动，苦干'),
                ('frozen', '/ˈfrəʊzn/', 'adj.', '冻结的，冰冻的'),
                ('rainy day', '/ˈreɪni deɪ/', 'n.', '雨天，困难时刻'),
            ],

            'grammar': [
                "Present Continuous for Future (现在进行时表将来): 'is hopping' — 描述正在发生的动作",
                "Present Continuous vs Present Simple (现在进行时与一般现在时对比): hopping vs passes",
                "Modal Verb 'shall' (shall的用法): 表示建议，在问句中较为正式",
                "Imperative Mood (祈使语气): 'Make sure you do not forget...' — 表示建议或命令",
                "Time Clauses with 'when' (when引导的时间状语从句): 'When winter came...'",
            ],

            'exercises': [
                "Compare the grasshopper's attitude with the ant's attitude. Which do you think is better? Why?",
                "What does 'a rainy day' mean in this context? Can you think of another expression with the same meaning?",
                "Find all the verbs in the present continuous tense and explain why they use this tense.",
                "Write a dialogue between the ant and grasshopper using at least 5 vocabulary words.",
                "Do you think the moral of this story is still relevant today? Write your opinion in English.",
            ]
        },

        {
            'title_en': 'The Lion and the Mouse',
            'title_zh': '狮子与老鼠',
            'level': 'Level 24',
            'text': """A lion was sleeping in his den when a little mouse began running up and down upon him. This awakened the lion.

The lion placed his huge paw upon the mouse and opened his big jaws to swallow him.

"Pardon, O King," cried the little mouse. "Forgive me this time, I shall never forget it. Who knows? Maybe I can help you someday."

The lion was so amused at the idea of the mouse helping him that he lifted his paw and let him go.

Some time later, a few hunters captured the lion. They tied him to a tree while they went to search for a wagon to carry him.

Just then the little mouse happened to pass by. Seeing the sad plight of the lion, he ran up to him and gnawed through the ropes. By and by, the lion was freed.

"Was I not right?" said the little mouse.

Little friends may prove to be great friends.""",

            'translation': """一只狮子正在洞里睡觉，一只小老鼠开始在它身上跑来跑去。这把狮子弄醒了。

狮子用巨大的爪子按住老鼠，张开大嘴要吞下它。

"饶命啊，大王，"小老鼠喊道。"请原谅我这一次，我永远不会忘记。也许有一天我能帮到您呢。"

狮子觉得老鼠要帮助他这个想法太有趣了，于是抬起爪子，放走了它。

过了一段时间，几个猎人捉住了狮子。他们把他绑在树上，然后去找车来把他运走。

就在那时，小老鼠恰好路过。看到狮子的悲惨处境，它跑过去咬断了绳子。不一会儿，狮子自由了。

"我说得对吧？"小老鼠说。

小朋友也可能成为大朋友。""",

            'vocab': [
                ('lion', '/ˈlaɪən/', 'n.', '狮子'),
                ('mouse', '/maʊs/', 'n.', '老鼠'),
                ('den', '/den/', 'n.', '兽穴，窝'),
                ('awaken', '/əˈweɪkən/', 'v.', '唤醒，使醒来'),
                ('paw', '/pɔː/', 'n.', '爪子'),
                ('jaws', '/dʒɔːz/', 'n.', '颌，嘴'),
                ('swallow', '/ˈswɒləʊ/', 'v.', '吞下，咽下'),
                ('pardon', '/ˈpɑːdn/', 'v./n.', '原谅，饶恕'),
                ('amused', '/əˈmjuːzd/', 'adj.', '觉得好笑的，被逗乐的'),
                ('gnaw', '/nɔː/', 'v.', '咬，啃'),
            ],

            'grammar': [
                "Past Continuous (过去进行时): 'was sleeping', 'was running' — 描述正在进行的过去动作",
                "Reported Speech (间接引语): 从 'Pardon, O King' 变为间接引语描述",
                "Adverbial Clause of Time (时间状语从句): 'Just then the little mouse happened to pass by'",
                "Modal Verb 'may' (may表示可能性): 'Little friends may prove to be great friends'",
                "Passive Voice (被动语态): 'the lion was freed', 'the lion was tied'",
            ],

            'exercises': [
                "Why did the lion let the mouse go? Do you think the lion was wise to do so?",
                "Identify all the passive voice sentences in the story and rewrite them in active voice.",
                "What does 'by and by' mean? Can you find it in the story and explain its usage?",
                "Write an alternative ending to this story where the mouse does not help the lion.",
                "Do you believe that small acts of kindness can lead to great rewards? Give an example.",
            ]
        },

        {
            'title_en': 'The Tortoise and the Hare',
            'title_zh': '乌龟与兔子',
            'level': 'Level 25',
            'text': """The hare was boasting about his speed one day when the tortoise said quietly: "Let us have a race."

The hare laughed at the idea. "You? Running against me? I could beat you while sleeping."

But the tortoise was not offended. He just smiled and said: "Try to beat me if you can."

The race began. The hare was so far ahead that he decided to take a nap under a tree. While the hare was sleeping, the tortoise kept walking slowly but steadily.

When the hare awoke, he thought the race was still far from over. He started running again as fast as he could. But by the time he reached the finish line, the tortoise was already there, smiling.

Slow and steady wins the race.

We can learn two lessons from this story. First, never underestimate your opponents. Second, steady progress can outshine flashy talent.""",

            'translation': """有一天，兔子正在吹嘘自己的速度，乌龟平静地说："我们比赛吧。"

兔子嘲笑这个主意。"你？跟我赛跑？我可以边睡觉边打败你。"

但乌龟并没有生气。他只是笑着说："如果可以的话，来打败我啊。"

比赛开始了。兔子遥遥领先，于是决定在一棵树下小睡一会儿。当兔子睡觉的时候，乌龟慢慢地但稳步地继续走着。

当兔子醒来时，它以为比赛还远远没有结束。它又开始尽可能快地跑起来。但当它到达终点线时，乌龟已经在那里，微笑着等待了。

慢而稳者获胜。

我们可以从这个故事中学到两课。第一，永远不要低估你的对手。第二，稳定的进步可以胜过华而不实的才能。""",

            'vocab': [
                ('tortoise', '/ˈtɔːtəs/', 'n.', '乌龟，陆龟'),
                ('hare', '/heə(r)/', 'n.', '野兔'),
                ('boasting', '/ˈbəʊstɪŋ/', 'n./v.', '吹嘘，自夸'),
                ('race', '/reɪs/', 'n./v.', '比赛，赛跑'),
                ('offended', '/əˈfendɪd/', 'adj.', '生气的，被冒犯的'),
                ('nap', '/næp/', 'n./v.', '小睡，打盹'),
                ('steadily', '/ˈstedəli/', 'adv.', '稳定地，稳步地'),
                ('finish line', '/ˈfɪnɪʃ laɪn/', 'n.', '终点线'),
                ('underestimate', '/ˌʌndərˈestɪmeɪt/', 'v.', '低估，看轻'),
                ('outshine', '/ˌaʊtˈʃaɪn/', 'v.', '胜过，超越'),
            ],

            'grammar': [
                'Reported Questions (间接疑问句): "...the tortoise said quietly: Let us have a race.",'
                "Past Simple vs Past Continuous (一般过去时与过去进行时对比): 用于描述同时发生的动作",
                "Causative 'let' (使役动词let): 'he decided to let him go' — 表示允许",
                "Modal Verb 'could' (could的用法): 'I could beat you while sleeping' — 表示能力",
                "Idiomatic Expression (习语): 'Slow and steady wins the race' — 经典谚语",
            ],

            'exercises': [
                "Why did the hare lose the race? What mistakes did the hare make?",
                "Do you agree with the moral 'Slow and steady wins the race'? Why or why not?",
                "Find three examples of the past simple tense and three examples of the past continuous tense in the story.",
                "Translate the second moral into Chinese: 'Steady progress can outshine flashy talent.'",
                "Write about a time when someone underestimated you but you proved them wrong. Use at least 5 vocabulary words from this lesson.",
            ]
        }
    ]

    # Process each story
    for story in stories:
        # Story header
        add_heading(doc, f"{story['level']}", level=2)
        add_heading(doc, f"{story['title_en']}", level=1)
        add_heading(doc, f"《{story['title_zh']}》", level=2)

        doc.add_paragraph()

        # Story content
        add_section_title(doc, "📖 Story Text 故事原文")

        # Split text into paragraphs and add bilingual
        paragraphs = story['text'].strip().split('\n\n')
        zh_paragraphs = story['translation'].strip().split('\n\n')

        for i, (en_p, zh_p) in enumerate(zip(paragraphs, zh_paragraphs)):
            add_bilingual_paragraph(doc, en_p.strip(), zh_p.strip())

        doc.add_paragraph()

        # Vocabulary
        add_section_title(doc, "📚 Vocabulary 词汇表")
        add_vocabulary_table(doc, story['vocab'])

        doc.add_paragraph()

        # Grammar
        add_section_title(doc, "📝 Grammar Points 语法点")
        add_grammar_section(doc, story['grammar'])

        doc.add_paragraph()

        # Exercises
        add_section_title(doc, "✏️ Practice Exercises 练习题")
        add_exercises(doc, story['exercises'])

        # Page break between stories
        doc.add_page_break()

    # ==================== END PAGE ====================
    add_heading(doc, "Congratulations! 恭喜完成 L21-25!", level=1)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🎉 You've completed all 5 Aesop's Fables! 🎉")
    set_font(run, 'Arial', 14, color=(0, 102, 51))

    doc.add_paragraph()

    summary = [
        "📖 Stories Mastered: 5",
        "📚 New Words Learned: 50",
        "✏️ Practice Exercises: 25",
        "🌟 Keep practicing and enjoy your English learning journey!"
    ]

    for item in summary:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_font(run, 'Arial', 12, color=(51, 51, 51))

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("英语学习 L21-25 | 伊索寓言进阶 | 2026")
    set_font(run, 'Arial', 10, color=(128, 128, 128))

    # Save document
    output_path = '/Users/mac/.openclaw/workspace/courses/英语学习_L21-25_AesopFables.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == '__main__':
    create_document()
    print("Done")
