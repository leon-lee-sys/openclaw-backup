#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
俄语水平测试 - 综合能力测试（30分钟完成）
"""
import os
import sys
sys.path.insert(0, '/opt/homebrew/lib/node_modules/openclaw/node_modules')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def set_style(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(30, 60, 114)
    return heading

def add_question(doc, num, question_text, options=None, is_translation=False):
    p = doc.add_paragraph()
    run = p.add_run(f"第{num}题：{question_text}")
    run.bold = True
    run.font.size = Pt(12)
    
    if options:
        for opt in options:
            p = doc.add_paragraph(opt, style='List Bullet')
    
    if is_translation:
        doc.add_paragraph("答案：________________")
        doc.add_paragraph()

def create_test():
    doc = Document()
    
    # Title
    title = doc.add_heading('俄语水平综合测试', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 60, 114)
    
    # Instructions
    p = doc.add_paragraph()
    p.add_run("考试时长：30分钟  |  总分：100分  |  考试形式：闭卷").font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # PART 1: 听力理解
    add_heading(doc, '第一部分：听力理解（共20分，每题4分）', 1)
    doc.add_paragraph("说明：阅读下列句子，选择正确的中文翻译。")
    
    add_question(doc, 1, "Как дела?", [
        "A. 你叫什么名字？",
        "B. 你怎么样？",
        "C. 你去哪里？",
        "D. 你在做什么？"
    ])
    
    add_question(doc, 2, "Где ты живёшь?", [
        "A. 你住在哪里？",
        "B. 你做什么工作？",
        "C. 你叫什么名字？",
        "D. 你从哪里来？"
    ])
    
    add_question(doc, 3, "Сколько времени?", [
        "A. 现在几点了？",
        "B. 今天星期几？",
        "C. 今天几号？",
        "D. 这个多少钱？"
    ])
    
    add_question(doc, 4, "Мне нужно идти.", [
        "A. 我想走了。",
        "B. 我需要走了。",
        "C. 我可以走了。",
        "D. 我必须留下。"
    ])
    
    add_question(doc, 5, "Это очень интересно.", [
        "A. 这太难了。",
        "B. 这很有趣。",
        "C. 这很无聊。",
        "D. 这很贵。"
    ])
    
    doc.add_page_break()
    
    # PART 2: 词汇与语法
    add_heading(doc, '第二部分：词汇与语法（共30分，每题3分）', 1)
    doc.add_paragraph("说明：选择最合适的选项完成句子。")
    
    add_question(doc, 6, "Я ___ студент.", [
        "A. есть",
        "B. это",
        "C. очень",
        "D. тоже"
    ])
    
    add_question(doc, 7, "___ пять часов.", [
        "A. Это",
        "B. Время",
        "C. Сейчас",
        "D. Уже"
    ])
    
    add_question(doc, 8, "Ты говоришь ___ русском?", [
        "A. на",
        "B. в",
        "C. с",
        "D. по"
    ])
    
    add_question(doc, 9, "Моя сестра ___ в Пекине.", [
        "A. живёт",
        "B. живёшь",
        "C. живём",
        "D. живут"
    ])
    
    add_question(doc, 10, "Вчера я ___ в театре.", [
        "A. буду",
        "B. был",
        "C. есть",
        "D. будет"
    ])
    
    add_question(doc, 11, "___ хороший фильм.", [
        "A. Этот",
        "B. Эта",
        "C. Это",
        "D. Эти"
    ])
    
    add_question(doc, 12, "Я люблю ___ музыку.", [
        "A. классический",
        "B. классическая",
        "C. классическое",
        "D. классическую"
    ])
    
    add_question(doc, 13, "Он ___ учитель.", [
        "A. русский",
        "B. русская",
        "C. русское",
        "D. русские"
    ])
    
    add_question(doc, 14, "Мы ___ друзья.", [
        "A. хороший",
        "B. хорошая",
        "C. хорошее",
        "D. хорошие"
    ])
    
    add_question(doc, 15, "___ книга интересная?", [
        "A. Какая",
        "B. Чья",
        "C. Эта",
        "D. Что"
    ])
    
    doc.add_page_break()
    
    # PART 3: 阅读理解
    add_heading(doc, '第三部分：阅读理解（共20分，每题5分）', 1)
    doc.add_paragraph("阅读下面的短文，回答问题。")
    
    doc.add_paragraph("Меня зовут Анна. Я из Москвы. Сейчас я учусь в университете на факультете русского языка. Мне очень нравится изучать иностранные языки. Я уже знаю английский и немецкий. Теперь я учу китайский язык. Китайский язык очень интересный, но очень трудный.")
    
    add_question(doc, 16, "Откуда Анна?", [
        "A. Из Пекина",
        "B. Из Москвы",
        "C. Из Киева",
        "D. Из Новосибирска"
    ])
    
    add_question(doc, 17, "На каком факультете учится Анна?", [
        "A. На историческом",
        "B. На факультете русского языка",
        "C. На медицинском",
        "D. На юридическом"
    ])
    
    add_question(doc, 18, "Какие языки уже знает Анна?", [
        "A. Китайский и японский",
        "B. Английский и немецкий",
        "C. Французский и испанский",
        "D. Итальянский и португальский"
    ])
    
    add_question(doc, 19, "Как Анна оценивает китайский язык?", [
        "A. Интересный, но лёгкий",
        "B. Скучный и трудный",
        "C. Интересный, но очень трудный",
        "D. Лёгкий и интересный"
    ])
    
    doc.add_page_break()
    
    # PART 4: 中译俄
    add_heading(doc, '第四部分：翻译（共15分，每题5分）', 1)
    doc.add_paragraph("将下列中文翻译成俄文。")
    
    add_question(doc, 20, "你好！你怎么样？", is_translation=True)
    add_question(doc, 21, "我叫李明。我来自上海。", is_translation=True)
    add_question(doc, 22, "今天天气很好。", is_translation=True)
    
    doc.add_page_break()
    
    # PART 5: 写作
    add_heading(doc, '第五部分：写作（共15分）', 1)
    doc.add_paragraph("用俄语写一篇短文，介绍一下自己（不少于5句话）。")
    doc.add_paragraph("内容包括：")
    doc.add_paragraph("• 你的名字", style='List Bullet')
    doc.add_paragraph("• 你来自哪里", style='List Bullet')
    doc.add_paragraph("• 你的工作或学习", style='List Bullet')
    doc.add_paragraph("• 你的爱好", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph("____________________________")
    doc.add_paragraph("____________________________")
    doc.add_paragraph("____________________________")
    doc.add_paragraph("____________________________")
    doc.add_paragraph("____________________________")
    doc.add_paragraph("____________________________")
    
    doc.add_page_break()
    
    # 参考答案
    add_heading(doc, '参考答案', 1)
    doc.add_paragraph("第一部分：听力理解")
    doc.add_paragraph("1-5: B A A B B")
    
    doc.add_paragraph()
    doc.add_paragraph("第二部分：词汇与语法")
    doc.add_paragraph("6-10: A C A A B")
    doc.add_paragraph("11-15: C D A D C")
    
    doc.add_paragraph()
    doc.add_paragraph("第三部分：阅读理解")
    doc.add_paragraph("16-19: B B B C")
    
    doc.add_paragraph()
    doc.add_paragraph("第四部分：翻译参考答案")
    doc.add_paragraph("20. Привет! Как дела? / Как ты?")
    doc.add_paragraph("21. Меня зовут Ли Мин. Я из Шанхая.")
    doc.add_paragraph("22. Сегодня хорошая погода.")
    
    doc.add_paragraph()
    doc.add_paragraph("第五部分：写作（参考范文）")
    doc.add_paragraph("Меня зовут Ли Мин. Я из Шанхая. Я работаю менеджером в компании. Мне нравится читать книги и слушать музыку. Я изучаю русский язык уже три месяца. Русский язык очень интересный!")
    
    # 保存
    output_path = "/Users/mac/.openclaw/workspace/courses/俄语水平测试.docx"
    doc.save(output_path)
    print(f"Done: {output_path}")
    return output_path

if __name__ == "__main__":
    create_test()
