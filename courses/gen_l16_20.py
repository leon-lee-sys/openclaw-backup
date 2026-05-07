#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英语学习L16-20进阶内容生成脚本
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUT_PATH = '/Users/mac/.openclaw/workspace/courses/英语学习_L16-20_AesopFables.docx'

doc = Document()

# 设置标题
title = doc.add_heading('Aesop\'s Fables — Lessons 16-20 (Advanced)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('高级进阶 · 中英文对照').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===== Lesson 16 =====
doc.add_heading('Lesson 16: The Fox and the Lion', 1)
p = doc.add_paragraph()
p.add_run('A Fox had never seen a Lion. When at last he came upon him, he was so terrified that he ran away and hid himself.').italic = True
doc.add_paragraph('狐狸从未见过狮子。当他终于遇到狮子时，他吓得逃跑了，躲藏了起来。')

doc.add_heading('📖 Vocabulary 词汇表', 2)
vocab = [
    ('Fox', 'n. 狐狸'),
    ('Lion', 'n. 狮子'),
    ('Terrified', 'adj. 极度恐惧的'),
    ('Hide', 'v. 躲藏'),
]
for w, m in vocab:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar 语法点', 2)
doc.add_paragraph('Past Simple 过去式: run → ran, hide → hid')

doc.add_heading('✏️ Exercise 练习', 2)
doc.add_paragraph('1. 用"terrified"造句')
doc.add_paragraph('2. 改写句子为过去式: "The Fox sees the Lion."')

# ===== Lesson 17 =====
doc.add_heading('Lesson 17: The Town Mouse and the Country Mouse', 1)
p = doc.add_paragraph()
p.add_run('A Town Mouse visited a Country Mouse. The Country Mouse offered simple food, but the Town Mouse longed for the rich food of the city.').italic = True
doc.add_paragraph('一只城里老鼠拜访了一只乡下老鼠。乡下老鼠用简单的食物款待他，但城里老鼠渴望城市的美味佳肴。')

doc.add_heading('📖 Vocabulary 词汇表', 2)
vocab = [
    ('Town Mouse', 'n. 城里老鼠'),
    ('Country Mouse', 'n. 乡下老鼠'),
    ('Simple', 'adj. 简单的'),
    ('Long for', 'v. 渴望'),
]
for w, m in vocab:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar 语法点', 2)
doc.add_paragraph('Comparative forms: simple → simpler, rich → richer')

doc.add_heading('✏️ Exercise 练习', 2)
doc.add_paragraph('1. 对比城里生活和乡下生活的优劣')
doc.add_paragraph('2. 用"long for"造句')

# ===== Lesson 18 =====
doc.add_heading('Lesson 18: The Lion, the Bear, and the Fox', 1)
p = doc.add_paragraph()
p.add_run('A Lion and a Bear were fighting over a piece of meat. A Fox passed by and offered to settle their dispute, only to take the meat for himself while they were distracted.').italic = True
doc.add_paragraph('狮子和熊为一块肉争吵。一只狐狸路过，提议为他们调解争端，却在他们分心时把肉拿走了。')

doc.add_heading('📖 Vocabulary 词汇表', 2)
vocab = [
    ('Fight over', 'v. 为...争吵'),
    ('Dispute', 'n. 争端'),
    ('Settle', 'v. 解决'),
    ('Distracted', 'adj. 分心的'),
]
for w, m in vocab:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar 语法点', 2)
doc.add_paragraph('"while" conjunction: 表示两个动作同时发生')

doc.add_heading('✏️ Exercise 练习', 2)
doc.add_paragraph('1. 这个故事告诉我们什么道理？')
doc.add_paragraph('2. 用"while"连接两个句子')

# ===== Lesson 19 =====
doc.add_heading('Lesson 19: The Fisherman and the Little Fish', 1)
p = doc.add_paragraph()
p.add_run('A Fisherman caught a tiny fish. The fish begged to be let go, promising to grow bigger and be more valuable later. The Fisherman refused, saying he would rather have one bird in hand than wait for many birds in the sky.').italic = True
doc.add_paragraph('渔夫捉到了一条小鱼。小鱼恳求被放走，承诺以后会长得更大更有价值。渔夫拒绝了，说他宁可要手中的一只鸟，也不要等天上的百鸟。')

doc.add_heading('📖 Vocabulary 词汇表', 2)
vocab = [
    ('Fisherman', 'n. 渔夫'),
    ('Tiny', 'adj. 微小的'),
    ('Beg', 'v. 恳求'),
    ('In hand', 'phrase. 手中的'),
]
for w, m in vocab:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar 语法点', 2)
doc.add_paragraph('Proverb 谚语: "A bird in hand is worth two in the bush." 手中的一鸟胜过灌木中的两鸟。')

doc.add_heading('✏️ Exercise 练习', 2)
doc.add_paragraph('1. 渔夫为什么拒绝放走小鱼？')
doc.add_paragraph('2. 讨论"知足常乐"的道理')

# ===== Lesson 20 =====
doc.add_heading('Lesson 20: The Wolf in Sheep\'s Clothing', 1)
p = doc.add_paragraph()
p.add_run('A Wolf decided to disguise himself in a sheep\'s skin to steal sheep. A Shepherd noticed the wolf but mistook him for a sheep. The wolf then attacked the flock.').italic = True
doc.add_paragraph('一只狼披上羊皮伪装自己来偷羊。牧羊人注意到了狼，但误以为他是羊。狼随后袭击了羊群。')

doc.add_heading('📖 Vocabulary 词汇表', 2)
vocab = [
    ('Disguise', 'v. 伪装'),
    ('Sheepskin', 'n. 羊皮'),
    ('Shepherd', 'n. 牧羊人'),
    ('Flock', 'n. 羊群'),
]
for w, m in vocab:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar 语法点', 2)
doc.add_paragraph('Idiom 成语: "Wolf in sheep\'s clothing" — 披着羊皮的狼，形容伪装成好人的坏人')

doc.add_heading('✏️ Exercise 练习', 2)
doc.add_paragraph('1. 这个故事适合用在哪里？（警惕伪装）')
doc.add_paragraph('2. 造句: "wolf in sheep\'s clothing"')

# 保存
doc.save(OUT_PATH)
print(f'Done: {OUT_PATH}')