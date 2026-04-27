#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英语学习 - W1D2: Simple Conversations (简单对话)
周二内容
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = '/Users/mac/.openclaw/workspace/courses/英语学习_W1D2.docx'

doc = Document()

# ===== 标题 =====
title = doc.add_heading('📚 Week 1: Beginner Basics', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_heading('Day 2: Simple Conversations（简单对话）', level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

level_note = doc.add_paragraph()
level_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = level_note.add_run('🎯 难度：初级 Beginner | ⏱ 建议学习时间：30分钟')
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ===== 今日主题 =====
doc.add_heading('🌟 今日主题：简单日常对话', level=2)

intro = doc.add_paragraph()
run = intro.add_run('学会这些对话，你就能和老外简单聊天了！')
run.italic = True

doc.add_paragraph()

# ===== Part 1: 对话1 =====
doc.add_heading('📖 Part 1: 对话1 — 问候', level=2)

dialogue1 = [
    ('A: Good morning! How are you?', 'A: 早上好！你好吗？'),
    ('B: Good morning! I\'m fine, thank you. And you?', 'B: 早上好！我很好，谢谢。你呢？'),
    ('A: I\'m doing well, thanks!', 'A: 我很好，谢谢！'),
    ('B: Nice to meet you!', 'B: 很高兴认识你！'),
    ('A: Nice to meet you too!', 'A: 我也很高兴认识你！'),
]

for en, cn in dialogue1:
    p = doc.add_paragraph()
    run = p.add_run(en)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'   中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 2: 对话2 =====
doc.add_heading('📖 Part 2: 对话2 — 介绍', level=2)

dialogue2 = [
    ('A: Hello! What\'s your name?', 'A: 你好！你叫什么名字？'),
    ('B: My name is Li. What\'s your name?', 'B: 我叫李。你叫什么名字？'),
    ('A: I\'m John. Where are you from?', 'A: 我是约翰。你来自哪里？'),
    ('B: I\'m from Shanghai. And you?', 'B: 我来自上海。你呢？'),
    ('A: I\'m from Beijing.', 'A: 我来自北京。'),
]

for en, cn in dialogue2:
    p = doc.add_paragraph()
    run = p.add_run(en)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'   中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 3: 对话3 =====
doc.add_heading('📖 Part 3: 对话3 — 告别', level=2)

dialogue3 = [
    ('A: Good night! See you tomorrow!', 'A: 晚安！明天见！'),
    ('B: Good night! See you!', 'B: 晚安！再见！'),
    ('A: Have a nice day!', 'A: 祝你有美好的一天！'),
    ('B: You too! Bye!', 'B: 你也是！拜拜！'),
]

for en, cn in dialogue3:
    p = doc.add_paragraph()
    run = p.add_run(en)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'   中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 4: 核心词汇 =====
doc.add_heading('📝 Part 4: 核心词汇', level=2)

words = [
    ('good morning', '-', '早上好', 'Good morning!'),
    ('good afternoon', '-', '下午好', 'Good afternoon!'),
    ('good evening', '-', '晚上好', 'Good evening!'),
    ('good night', '-', '晚安（告别时）', 'Good night! See you tomorrow!'),
    ('fine', '/faɪn/', '好的、不错的', 'I\'m fine, thank you.'),
    ('nice to meet you', '-', '很高兴认识你', 'Nice to meet you!'),
    ('where', '/weə/', '在哪里', 'Where are you from?'),
    ('from', '/frɒm/', '来自', 'I\'m from Shanghai.'),
]

for word, phonetic, meaning, example in words:
    p = doc.add_paragraph()
    run = p.add_run(f'• {word} ')
    run.bold = True
    run.font.color.rgb = RGBColor(30, 60, 114)
    run = p.add_run(f'— {meaning}')
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  例句：{example}')
    run2.italic = True
    run2.font.size = Pt(10)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 5: 语法 =====
doc.add_heading('🔧 Part 5: 语法 — be动词', level=2)

p = doc.add_paragraph()
run = p.add_run('be动词（am/is/are）= 中文的"是"')
run.bold = True
run.font.color.rgb = RGBColor(30, 60, 114)

doc.add_paragraph()

grammar = [
    ('I', 'am (I\'m)', '我 是'),
    ('You', 'are (You\'re)', '你 是'),
    ('He/She/It', 'is (He\'s/She\'s/It\'s)', '他/她/它 是'),
    ('We/They', 'are (We\'re/They\'re)', '我们/他们 是'),
]

for subject, verb, cn in grammar:
    p = doc.add_paragraph()
    run = p.add_run(f'• {subject} {verb}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'   中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== Part 6: 练习 =====
doc.add_heading('✍️ Part 6: 今日口语练习', level=2)

p = doc.add_paragraph()
run = p.add_run('任务：用英语和别人打招呼并介绍自己（至少3句）')
run.bold = True

doc.add_paragraph()

practice = [
    ('Hello! How are you?', '你好！你好吗？'),
    ('I\'m fine, thank you.', '我很好，谢谢。'),
    ('My name is [你的名字].', '我叫[你的名字]。'),
    ('I\'m from [你的城市].', '我来自[你的城市]。'),
    ('Nice to meet you!', '很高兴认识你！'),
    ('See you later!', '回头见！'),
]

for en, cn in practice:
    p = doc.add_paragraph()
    run = p.add_run(f'• {en}')
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'   中文：{cn}')
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ===== 今日总结 =====
doc.add_heading('📋 今日总结', level=2)

summary = [
    ('对话', '3个简单对话场景（问候/介绍/告别）'),
    ('词汇', '8个常用词汇'),
    ('语法', 'be动词 am/is/are 的用法'),
    ('练习', '用英语打招呼和自我介绍'),
]

for item, desc in summary:
    p = doc.add_paragraph()
    run = p.add_run(f'✅ {item}：')
    run.bold = True
    run = p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('🌟 Week 1 - Day 2 | 初级难度 | 每天半小时，坚持就是胜利！')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUT)
print(f"Saved: {OUT}")
