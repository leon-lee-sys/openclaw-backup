#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英语学习L21-25进阶内容生成脚本
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = '/Users/mac/.openclaw/workspace/courses/英语学习_L21-25_AesopFables.docx'

doc = Document()

title = doc.add_heading('Aesop\'s Fables — Lessons 21-25 (Advanced)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('高级进阶 · 中英文对照').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===== Lesson 21 =====
doc.add_heading('Lesson 21: The Crow and the Pitcher', 1)
p = doc.add_paragraph()
p.add_run('A Crow, half-dead with thirst, saw a pitcher with some water in it. The Crow tried to drink but found the neck of the pitcher too narrow. At last he hit upon a clever plan.').italic = True
doc.add_paragraph('一只乌鸦渴得半死，看见一只 pitcher（细口瓶）里有水。它想喝但瓶口太窄。最后它想出了一个妙计。')

doc.add_heading('📖 Vocabulary', 2)
for w, m in [('Crow', 'n. 乌鸦'), ('Pitcher', 'n. 细口瓶'), ('Thirst', 'n. 口渴'), ('Clever', 'adj. 聪明的')]:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar', 2)
doc.add_paragraph('Past Simple: saw → see, found → find, hit → hit')

doc.add_heading('✏️ Exercise', 2)
doc.add_paragraph('1. What was the clever plan?')
doc.add_paragraph('2. 翻译：半死')

# ===== Lesson 22 =====
doc.add_heading('Lesson 22: The Fox and the Grapes', 1)
p = doc.add_paragraph()
p.add_run('A Fox one hot day saw a bunch of grapes hanging from a vine. He tried to reach them but found they were too high. He walked away saying, "I\'m sure they are sour."').italic = True
doc.add_paragraph('一天，狐狸在大热天看见一串葡萄挂在藤上。它想吃但发现太高了。它走开说："我确定它们是酸的。"

doc.add_heading('📖 Vocabulary', 2)
for w, m in [('Grapes', 'n. 葡萄'), ('Vine', 'n. 藤'), ('Sour', 'adj. 酸的'), ('Reach', 'v. 够到')]:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Idiom', 2)
doc.add_paragraph('"Sour grapes" — 吃不到葡萄说葡萄酸')

doc.add_heading('✏️ Exercise', 2)
doc.add_paragraph('1. 用"sour grapes"造句')

# ===== Lesson 23 =====
doc.add_heading('Lesson 23: The Ant and the Grasshopper', 1)
p = doc.add_paragraph()
p.add_run('In a field one summer\'s day, a Grasshopper was hopping and singing. An Ant passed by with grain to store for winter. "Why not come and chat with me?" asked the Grasshopper. The Ant replied, "I am storing food for winter."').italic = True
doc.add_paragraph('夏天，一只蚱蜢在田野里跳来跳去唱歌。一只蚂蚁背着谷物走过，要为冬天储存食物。"为什么不来和我聊天？"蚱蜢问。蚂蚁回答："我在为冬天储存食物。"

doc.add_heading('📖 Vocabulary', 2)
for w, m in [('Grasshopper', 'n. 蚱蜢'), ('Ant', 'n. 蚂蚁'), ('Hop', 'v. 跳跃'), ('Grain', 'n. 谷物')]:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar', 2)
doc.add_paragraph('Present Continuous vs Future: "I am storing" vs "I will store"')

doc.add_heading('✏️ Exercise', 2)
doc.add_paragraph('1. 这个故事的寓意是什么？')

# ===== Lesson 24 =====
doc.add_heading('Lesson 24: The Lion and the Mouse', 1)
p = doc.add_paragraph()
p.add_run('A Lion was sleeping when a Mouse began to run up and down upon him. The Lion arose and shook his mighty paw. A tiny Mouse answered, "Pardon me, and I will repay your kindness."').italic = True
doc.add_paragraph('狮子正在睡觉，一只老鼠开始在它身上跑来跑去。狮子站起来挥动它有力的爪子。一只小老鼠说："请原谅我，我会报答您的好意。"

doc.add_heading('📖 Vocabulary', 2)
for w, m in [('Mighty', 'adj. 强大的'), ('Paw', 'n. 爪子'), ('Repay', 'v. 报答'), ('Kindness', 'n. 仁慈')]:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Grammar', 2)
doc.add_paragraph('"Will" vs "Would" in polite requests')

doc.add_heading('✏️ Exercise', 2)
doc.add_paragraph('1. 故事的结局是什么？')

# ===== Lesson 25 =====
doc.add_heading('Lesson 25: The Tortoise and the Hare', 1)
p = doc.add_paragraph()
p.add_run('A Hare ridiculed a Tortoise for moving so slowly. The Tortoise replied, "Give me a head start and I will beat you in a race." The Hare agreed and took a nap. Meanwhile, the Tortoise kept walking. When the Hare woke up, the Tortoise had already won.').italic = True
doc.add_paragraph('一只野兔嘲笑乌龟走得慢。乌龟说："让我先走一段，我会赛过你。"兔子同意了，然后打了个盹。同时，乌龟一直走。当兔子醒来时，乌龟已经赢了。"

doc.add_heading('📖 Vocabulary', 2)
for w, m in [('Tortoise', 'n. 乌龟'), ('Hare', 'n. 野兔'), ('Ridiculed', 'v. 嘲笑'), ('Beat', 'v. 击败')]:
    doc.add_paragraph(f'• {w} — {m}')

doc.add_heading('📝 Idiom', 2)
doc.add_paragraph('"Slow and steady wins the race." — 稳扎稳打才能赢')

doc.add_heading('✏️ Exercise', 2)
doc.add_paragraph('1. 这个故事告诉我们什么道理？')

doc.save(OUT_PATH)
print(f'Done: {OUT_PATH}')