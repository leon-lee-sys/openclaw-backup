from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from bs4 import BeautifulSoup

doc = Document()

# Set default font for Chinese
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# Title
title = doc.add_heading('英语学习材料 · 初级第1-5课', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('📚 学习计划：先初级，再中级，循序渐进 🐦')
doc.add_paragraph('')

# Lesson 1: The Boy Who Cried Wolf (Aesop's Fables)
doc.add_heading('Lesson 1: The Boy Who Cried Wolf', level=2)
doc.add_paragraph('伊索寓言 · 狼来了')

original1 = '''A shepherd boy watched his flock near a village. He was often bored and wanted to have some fun.

"Help! Wolf! Wolf!" he shouted.

The villagers came running. But when they arrived, they found no wolf. The boy laughed at them.

He did this many times. The villagers grew angry.

Then one day, a wolf really came. The boy cried out loudly: "Help! Wolf! Wolf!"

But this time, no one came. The villagers thought he was joking again.

The wolf scattered the flock and ran away.'''

doc.add_paragraph(original1)
doc.add_paragraph('')

vocab1 = '''核心词汇：
• shepherd /ˈʃepərd/ n. 牧羊人
• flock /flɒk/ n. 羊群
• village /ˈvɪlɪdʒ/ n. 村庄
• bored /bɔːrd/ adj. 无聊的
• fun /fʌn/ n. 乐趣
• shout /ʃaʊt/ v. 大喊
• villager /ˈvɪlɪdʒər/ n. 村民
• laugh /læf/ v. 笑
• angry /ˈæŋɡri/ adj. 生气的
• really /ˈrɪəli/ adv. 真的
• scatter /ˈskætər/ v. 驱散'''

doc.add_paragraph(vocab1)

grammar1 = '''语法要点：
1. 一般过去时 (Past Simple): watched, shouted, came, found, did
   构成：主语 + 动词过去式
   例：He watched his flock. / The wolf scattered the flock.

2. 过去进行时 (Past Continuous): was watching, were running
   构成：主语 + was/were + 动词-ing
   例：He was often bored. / The villagers were laughing.

3. 间接引语：The villagers thought he was joking.
   动词变化：comes → came (过去式)'''

doc.add_paragraph(grammar1)
doc.add_paragraph('')

# Lesson 2: The Tortoise and the Hare
doc.add_heading('Lesson 2: The Tortoise and the Hare', level=2)
doc.add_paragraph('伊索寓言 · 龟兔赛跑')

original2 = '''A hare was very proud of his speed. He laughed at a tortoise.

"Let's race!" said the hare. "You'll never catch me."

The tortoise agreed. They started at the same time.

The hare ran fast at first. Then he thought he had enough time. He took a nap under a tree.

The tortoise walked slowly but never stopped. He passed the sleeping hare and won the race.

"Slow and steady wins the race," said the tortoise.'''

doc.add_paragraph(original2)
doc.add_paragraph('')

vocab2 = '''核心词汇：
• hare /her/ n. 野兔
• proud /praʊd/ adj. 骄傲的
• speed /spiːd/ n. 速度
• tortoise /ˈtɔːrtəs/ n. 乌龟
• laugh at /læf æt/ 嘲笑
• race /reɪs/ n./v. 赛跑
• never /ˈnevər/ adv. 从不
• catch /kætʃ/ v. 赶上
• start /stɑːrt/ v. 开始
• nap /næp/ n. 小睡
• slowly /ˈsloʊli/ adv. 慢慢地
• steady /ˈstedi/ adj. 稳定的'''

doc.add_paragraph(vocab2)

grammar2 = '''语法要点：
1. 一般将来时 (Future Simple): will + 动词原形
   例：You'll never catch me. / He will take a nap.

2. 形容词比较级：faster, slower
   构成：-er / more + 形容词
   例：The tortoise walked slower than the hare.

3. 情态动词：can, will, never
   例：You'll never catch me. / He can run fast.'''

doc.add_paragraph(grammar2)
doc.add_paragraph('')

# Lesson 3: The Lion and the Mouse
doc.add_heading('Lesson 3: The Lion and the Mouse', level=2)
doc.add_paragraph('伊索寓言 · 狮子与老鼠')

original3 = '''A lion was sleeping. A little mouse walked over him by accident. The lion woke up and caught the mouse.

"Please let me go!" said the mouse. "If you spare me, I will help you someday."

The lion laughed. "How can a tiny mouse help a lion?" But he let the mouse go.

Days later, some hunters caught the lion. They tied him to a tree with a rope.

The mouse heard the lion's roar. He ran to help. He gnawed the rope with his sharp teeth.

Soon the lion was free. "Thank you, little mouse," said the lion.

"Sometimes the small can help the big," said the mouse.'''

doc.add_paragraph(original3)
doc.add_paragraph('')

vocab3 = '''核心词汇：
• lion /ˈlaɪən/ n. 狮子
• mouse /maʊs/ n. 老鼠
• sleep /sliːp/ v. 睡觉
• awake /əˈweɪk/ v. 醒来
• catch /kætʃ/ v. 抓住
• spare /sper/ v. 放过
• tiny /ˈtaɪni/ adj. 微小的
• hunter /ˈhʌntər/ n. 猎人
• tie /taɪ/ v. 捆绑
• rope /roʊp/ n. 绳子
• gnaw /nɔː/ v. 咬
• sharp /ʃɑːrp/ adj. 锋利的'''

doc.add_paragraph(vocab3)

grammar3 = '''语法要点：
1. 过去式规则变化：
   动名词+ed: walked, laughed, helped
   不规则: slept → slept, woke → woke, caught → caught

2. 条件句 (Type 2): If + past tense + would
   例：If you spare me, I will help you.
   (If you spared me, I would help you.) - 与现在事实相反

3. 形容词最高级：the smallest, the fastest
   the + 形容词最高级 + in/of 介词短语'''

doc.add_paragraph(grammar3)
doc.add_paragraph('')

# Lesson 4: The Fox and the Grapes
doc.add_heading('Lesson 4: The Fox and the Grapes', level=2)
doc.add_paragraph('伊索寓言 · 狐狸与葡萄')

original4 = '''A fox was very hungry. He saw some sweet grapes hanging from a vine.

He wanted them very much. He jumped and jumped, but he could not reach them.

After many tries, he gave up. He walked away saying:

"Those grapes are probably sour anyway. I didn't really want them."

Sometimes we pretend we don't want what we cannot have.'''

doc.add_paragraph(original4)
doc.add_paragraph('')

vocab4 = '''核心词汇：
• fox /fɒks/ n. 狐狸
• hungry /ˈhʌŋɡri/ adj. 饥饿的
• sweet /swiːt/ adj. 甜的
• grape /ɡreɪp/ n. 葡萄
• hang /hæŋ/ v. 悬挂
• vine /vaɪn/ n. 葡萄藤
• reach /riːtʃ/ v. 够到
• try /traɪ/ n./v. 尝试
• give up /ɡɪv ʌp/ 放弃
• probably /ˈprɒbəbli/ adv. 可能
• sour /ˈsaʊər/ adj. 酸的
• anyway /ˈeniweɪ/ adv. 反正
• pretend /prɪˈtend/ v. 假装'''

doc.add_paragraph(vocab4)

grammar4 = '''语法要点：
1. 情态动词：could, can, may, might
   例：He could not reach them. / I may be wrong.

2. 不定式：to + 动词原形
   例：He wanted to eat. / She tried to jump.

3. 动名词：动词-ing作主语或宾语
   例：Jumping is fun. / He gave up trying.'''

doc.add_paragraph(grammar4)
doc.add_paragraph('')

# Lesson 5: The Ugly Duckling (Simplified Hans Christian Andersen)
doc.add_heading('Lesson 5: The Ugly Duckling', level=2)
doc.add_paragraph('安徒生童话 · 丑小鸭')

original5 = '''A mother duck had five eggs. Four hatched into beautiful yellow ducklings. The fifth egg was very large.

The last egg finally cracked. Out came a big, gray duckling. He was ugly.

"Your brother is so ugly!" said the other birds. They were unkind to him.

The ugly duckling felt very sad. He left home and walked alone.

Winter came. The pond froze. A kind farmer found him and took him home.

Spring arrived. The ugly duckling grew big and strong. He looked at his reflection. He was now a beautiful swan!

All the other swans swam to greet him. He was no longer sad.'''

doc.add_paragraph(original5)
doc.add_paragraph('')

vocab5 = '''核心词汇：
• duck /dʌk/ n. 鸭子
• egg /eɡ/ n. 蛋
• hatch /hætʃ/ v. 孵化
• beautiful /ˈbjuːtɪfl/ adj. 美丽的
• yellow /ˈjeloʊ/ adj. 黄色的
• ugly /ˈʌɡli/ adj. 丑陋的
• duckling /ˈdʌklɪŋ/ n. 小鸭子
• unkind /ʌnˈkaɪnd/ adj. 不友善的
• leave /liːv/ v. 离开
• alone /əˈloʊn/ adj./adv. 独自的
• freeze /friːz/ v. 结冰
• farmer /ˈfɑːrmər/ n. 农夫
• spring /sprɪŋ/ n. 春天
• reflection /rɪˈplekʃn/ n. 倒影
• swan /swɒn/ n. 天鹅'''

doc.add_paragraph(vocab5)

grammar5 = '''语法要点：
1. 一般过去时 vs. 过去进行时：
   过去进行时强调正在发生：The other birds were unkind to him.
   一般过去时描述已完成的动作：He left home.

2. 形容词比较级和最高级：
   ugly → uglier → ugliest
   beautiful → more beautiful → most beautiful
   good → better → best

3. 连词：but, and, so, because
   He was ugly, but he was kind. / He left because he was sad.'''

doc.add_paragraph(grammar5)
doc.add_paragraph('')

# Summary page
doc.add_heading('📝 每日学习指南', level=2)

guide = '''学习步骤：
1. 先通读全文，理解故事内容
2. 划出不认识的单词，查词典
3. 重读课文，尝试翻译成中文
4. 背诵核心词汇（建议用词卡）
5. 用重点语法造句，练习写作
6. 朗读课文，注意发音和语感

每日任务：
• 朗读2-3遍
• 背诵5-8个核心词汇
• 用今天学的语法造3个句子
• 把你的朗读录音发给我，我帮你纠正发音 🐦

坚持下去，积少成多！'''

doc.add_paragraph(guide)

# Save
doc.save('/Users/mac/Desktop/英语学习_初级_Lessons_1-5.docx')
print('Document created successfully!')
