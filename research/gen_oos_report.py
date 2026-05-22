#!/usr/bin/env python3
"""生成商业在轨维护与服务研究报告"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

doc = Document()

# 标题
title = doc.add_heading('商业在轨维护与服务（On-Orbit Servicing）', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('研究报告')
doc.add_paragraph(f'报告日期：{date.today().strftime("%Y年%m月%d日")}')
doc.add_paragraph('研究机构：小燕子研究院')
doc.add_paragraph('')

# ============ 第一章 行业概述 ============
doc.add_heading('一、行业概述', 1)

doc.add_heading('1.1 行业定义', 2)
doc.add_paragraph(
    '在轨维护与服务（On-Orbit Servicing, OOS）是指在太空中对卫星、空间站等航天器进行的'
    '维修、延寿、加油、碎片清除、轨道调整等一系列商业化服务活动。'
    '这一新兴行业正在从根本上改变太空经济的商业逻辑——从"一次性使用"转向"持续运维"。'
)

doc.add_heading('1.2 行业发展背景', 2)
p = doc.add_paragraph()
p.add_run('四大驱动力：').bold = True
doc.add_paragraph('① 卫星寿命限制：地球静止轨道（GEO）卫星燃料耗尽是主要失效模式，在轨加油可延长卫星寿命5-15年', style='List Bullet')
doc.add_paragraph('② 空间碎片威胁：目前约有9,000吨太空碎片，在轨服务是唯一主动清除手段', style='List Bullet')
doc.add_paragraph('③ 军用价值：美国太空军视"动态太空操作"为核心能力，可在战时机动卫星规避威胁或接近对手卫星', style='List Bullet')
doc.add_paragraph('④ 成本压力：GEO卫星造价动辄5-10亿美元，延寿服务性价比极高', style='List Bullet')

doc.add_heading('1.3 市场规模', 2)
# 市场规模表格
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '指标'
hdr[1].text = '数据'
hdr[2].text = '来源'
data = [
    ('2025年全球市场规模', '约46.7亿美元', 'Precedence Research'),
    ('2026年预测', '约51.6亿美元', 'Precedence Research'),
    ('2035年预测', '约126亿美元', 'Precedence Research（10.43% CAGR）'),
    ('2030年预测', '约61.4亿美元', 'Spherical Insights（10.7% CAGR）'),
]
for i, (a, b, c) in enumerate(data, 1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
    table.rows[i].cells[2].text = c

doc.add_paragraph('')
doc.add_paragraph('北美市场占全球份额约37%（2025年），欧洲市场预计2030年超50亿美元（年增速约11.5%）。')

# ============ 第二章 国际主要公司 ============
doc.add_heading('二、国际主要公司分析', 1)

# 2.1 诺斯罗普·格鲁曼
doc.add_heading('2.1 诺斯罗普·格鲁曼（Northrop Grumman）', 2)
doc.add_paragraph('基本情况：美国防务巨头，在轨服务领域先驱企业')
doc.add_paragraph('核心业务：卫星延寿、轨道转移、军事卫星维修', style='List Bullet')
doc.add_paragraph('代表项目：', style='List Bullet')
doc.add_paragraph('  • MEV-1（任务扩展飞行器1号）：2019年发射，为Intelsat 901卫星提供延寿服务，在轨运行至今', style='List Bullet')
doc.add_paragraph('  • MEV-2：2021年发射，为Intelsat 1002提供延寿服务', style='List Bullet')
doc.add_paragraph('  • Mission Extension Pod（MEP）：与Astroscale合作开发的小型推进飞行器', style='List Bullet')
doc.add_paragraph('优势：深厚军方关系、成熟的空间系统集成能力', style='List Bullet')
doc.add_paragraph('劣势：商业模式以政府项目为主，商业化程度有限', style='List Bullet')

# 2.2 Maxar Technologies
doc.add_heading('2.2 Maxar Technologies（美国）', 2)
doc.add_paragraph('基本情况：全球领先的空间技术公司，旗下拥有DigitalGlobe等卫星影像业务')
doc.add_paragraph('核心业务：卫星影像、太空机器人、在轨服务', style='List Bullet')
doc.add_paragraph('代表项目：Space Infrastructure Dexterous Robot（SPIDER）：太空基础设施灵巧机器人，可执行卫星维修与在轨装配', style='List Bullet')
doc.add_paragraph('优势：卫星影像数据支撑、空间机器人技术积累', style='List Bullet')

# 2.3 Astroscale
doc.add_heading('2.3 Astroscale Holdings（日本）', 2)
doc.add_paragraph('基本情况：全球首家商业在轨服务上市公司，总部东京，2021年登陆东京证券交易所')
doc.add_paragraph('核心业务：主动碎片清除（ADR）、卫星延寿、在轨检查', style='List Bullet')
doc.add_paragraph('代表项目：', style='List Bullet')
doc.add_paragraph('  • ADRAS-J：全球首次商业化主动碎片清除演示（2024年）', style='List Bullet')
doc.add_paragraph('  • CLASP：与欧洲航天局（ESA）合作项目', style='List Bullet')
doc.add_paragraph('  • ELSA-d：在轨碎片清除演示（2021年）', style='List Bullet')
doc.add_paragraph('优势：上市公司、全球唯一已验证商业ADR技术、日本政府背书', style='List Bullet')
doc.add_paragraph('劣势：技术仍处于早期商业化阶段，盈利能力待验证', style='List Bullet')

# 2.4 SpaceX
doc.add_heading('2.4 SpaceX（美国）', 2)
doc.add_paragraph('基本情况：全球最大私营航天企业，估值约3,500亿美元')
doc.add_paragraph('核心业务：火箭发射、星链卫星、在轨服务技术储备', style='List Bullet')
doc.add_paragraph('代表能力：', style='List Bullet')
doc.add_paragraph('  • Starship设计可用于大型在轨任务', style='List Bullet')
doc.add_paragraph('  • 星舰在轨加油技术（2026年计划验证）', style='List Bullet')
doc.add_paragraph('  • Dragon飞船具备对接能力，理论可扩展至在轨维修', style='List Bullet')
doc.add_paragraph('优势：运载能力独步全球、成本控制能力强', style='List Bullet')
doc.add_paragraph('劣势：战略重心在星链和深空探测，在轨服务非核心业务', style='List Bullet')

# 2.5 法国SpaceNEO
doc.add_heading('2.5 SpaceNEO（欧洲）', 2)
doc.add_paragraph('基本情况：法国初创企业，专注GEO卫星延寿服务，由欧洲航天局支持')
doc.add_paragraph('核心业务：GEO卫星延寿飞行器', style='List Bullet')
doc.add_paragraph('目标：2028年前具备商业服务能力', style='List Bullet')

# 2.6 印度
doc.add_heading('2.6 印度（国际合作视角）', 2)
doc.add_paragraph('印度空间研究组织（ISRO）积极布局在轨服务：', style='List Bullet')
doc.add_paragraph('  • 2022年完成"太空对接"实验（Spadex）', style='List Bullet')
doc.add_paragraph('  • 规划2030年前具备卫星延寿能力', style='List Bullet')
doc.add_paragraph('印度凭借成本优势可能成为未来重要竞争者', style='List Bullet')

# ============ 第三章 国内主要公司 ============
doc.add_heading('三、国内主要公司分析', 1)

doc.add_heading('3.1 中国航天科技集团（CASC）', 2)
doc.add_paragraph('基本情况：中国航天工业旗舰，全球最大航天企业之一')
doc.add_paragraph('核心业务：运载火箭、卫星、空间站、在轨服务技术研发', style='List Bullet')
doc.add_paragraph('重大进展：', style='List Bullet')
doc.add_paragraph('  • 2024年6月：实践-19号卫星成功进行世界首次GEO在轨加油演示', style='List Bullet')
doc.add_paragraph('  • 实践-21号：实验性GEO碎片减缓技术', style='List Bullet')
doc.add_paragraph('  • 多家研究所从事机械臂、太空机器人研发', style='List Bullet')
doc.add_paragraph('特点：国家战略驱动，技术积累深厚，但商业化程度较低', style='List Bullet')

doc.add_heading('3.2 银河航天', 2)
doc.add_paragraph('基本情况：中国商业航天领军企业，专注低轨宽带互联网卫星')
doc.add_paragraph('核心业务：银河Galaxy卫星星座建设、卫星平台技术', style='List Bullet')
doc.add_paragraph('在轨服务布局：', style='List Bullet')
doc.add_paragraph('  • 规划卫星延寿服务能力', style='List Bullet')
doc.add_paragraph('  • 低轨星座维护需求催生在轨服务市场', style='List Bullet')

doc.add_heading('3.3 长光卫星技术股份有限公司', 2)
doc.add_paragraph('基本情况：科创板上市，商业遥感卫星企业龙头')
doc.add_paragraph('核心业务：遥感卫星研制与运营、"吉林一号"星座', style='List Bullet')
doc.add_paragraph('在轨服务潜力：遥感卫星寿命延长服务需求广阔', style='List Bullet')

doc.add_heading('3.4 蓝箭航天空间科技', 2)
doc.add_paragraph('基本情况：商业运载火箭企业，研朱雀二号可复用火箭')
doc.add_paragraph('在轨服务布局：规划火箭子级在轨回收复用', style='List Bullet')

doc.add_heading('3.5 国内行业整体评估', 2)
p = doc.add_paragraph()
p.add_run('优势：').bold = True
doc.add_paragraph('① 2024年GEO在轨加油技术世界领先', style='List Bullet')
doc.add_paragraph('② 国家战略高度重视，政策支持明确', style='List Bullet')
doc.add_paragraph('③ 完整的航天工业体系支撑', style='List Bullet')
p = doc.add_paragraph()
p.add_run('劣势：').bold = True
doc.add_paragraph('① 商业化程度低，主要依赖国家任务', style='List Bullet')
doc.add_paragraph('② 上市公司少，投资标的稀缺', style='List Bullet')
doc.add_paragraph('③ 核心技术仍与美国有差距', style='List Bullet')

# ============ 第四章 关键技术 ============
doc.add_heading('四、关键技术分类', 1)

doc.add_heading('4.1 卫星延寿与服务', 2)
doc.add_paragraph('定义：为燃料耗尽的卫星提供推进剂补给或轨道维持服务')
doc.add_paragraph('代表公司：诺斯罗普·格鲁曼（MEV系列）、SpaceNEO', style='List Bullet')
doc.add_paragraph('技术难点：远距离逼近、GEO高轨转移、对接稳定性', style='List Bullet')

doc.add_heading('4.2 轨道转移与在轨加油', 2)
doc.add_paragraph('定义：将飞行器从一个轨道转移到另一个，并进行流体（推进剂）转移')
doc.add_paragraph('代表项目：NASA太空机器人项目（Space Robotics），中国实践-19号', style='List Bullet')
doc.add_paragraph('技术难点：轨道力学计算、低温推进剂管理、对接密封', style='List Bullet')

doc.add_heading('4.3 主动碎片清除（ADR）', 2)
doc.add_paragraph('定义：主动捕获并清除失效卫星、火箭残骸等空间碎片')
doc.add_paragraph('代表公司：Astroscale（全球领先）、Tethers Unlimited', style='List Bullet')
doc.add_paragraph('技术难点：非合作目标捕获、碎片追踪、规避碰撞', style='List Bullet')

doc.add_heading('4.4 在轨装配与制造', 2)
doc.add_paragraph('定义：在太空中直接组装大型结构或制造复杂部件')
doc.add_paragraph('代表项目：DARPA的RSGS项目（机器人卫星加油与服务）、Maxar的SPIDER', style='List Bullet')
doc.add_paragraph('未来愿景：建设"月球门户"等大型深空基础设施', style='List Bullet')

doc.add_heading('4.5 机械臂与太空机器人', 2)
doc.add_paragraph('定义：在轨操作的核心使能技术，负责捕获、检查、维修等任务')
doc.add_paragraph('代表公司：Maxar、Tethers Unlimited、NASA', style='List Bullet')
doc.add_paragraph('技术趋势：AI驱动的自主操作、触觉反馈远程操控', style='List Bullet')

# ============ 第五章 市场前景 ============
doc.add_heading('五、市场前景与投资机会', 1)

doc.add_heading('5.1 市场驱动因素', 2)
doc.add_paragraph('① GEO卫星换新成本高：单星5-10亿美元，延寿仅需数千万美元', style='List Bullet')
doc.add_paragraph('② 太空军备竞赛：中美俄在轨对抗能力建设催生大量需求', style='List Bullet')
doc.add_paragraph('③ ESG投资热潮：空间碎片清理符合可持续发展理念', style='List Bullet')
doc.add_paragraph('④ 卫星星座爆发：星链、一网等万颗级星座急需运维手段', style='List Bullet')

doc.add_heading('5.2 投资机会', 2)
p = doc.add_paragraph()
p.add_run('一级市场：').bold = True
doc.add_paragraph('• Astroscale（东京证交所上市，代码186A）：全球唯一纯正OOS标的', style='List Bullet')
doc.add_paragraph('• SpaceNEO（法国）：欧洲OOS龙头，IPO筹备中', style='List Bullet')
doc.add_paragraph('• 各防务巨头的太空服务部门（诺格、洛马、雷神）', style='List Bullet')
p = doc.add_paragraph()
p.add_run('二级市场：').bold = True
doc.add_paragraph('• 诺斯罗普·格鲁曼（NYSE：NOC）：MEV项目稳定运营', style='List Bullet')
doc.add_paragraph('• Maxar（NYSE：MAXR）：技术储备丰富但财务压力大', style='List Bullet')
doc.add_paragraph('• SpaceX（私营）：估值已高但成长空间仍大', style='List Bullet')
doc.add_paragraph('• A股：中国卫星网络集团（筹备中）、航天科技集团旗下上市公司', style='List Bullet')

doc.add_heading('5.3 风险因素', 2)
doc.add_paragraph('① 技术风险：太空环境恶劣，任务失败率仍较高', style='List Bullet')
doc.add_paragraph('② 地缘政治：军事应用敏感性可能导致出口管制', style='List Bullet')
doc.add_paragraph('③ 商业模式待验证：大多数OOS公司尚未实现盈利', style='List Bullet')
doc.add_paragraph('④ 标准化缺失：不同卫星接口不统一，制约商业服务规模化', style='List Bullet')

# ============ 第六章 结论 ============
doc.add_heading('六、结论', 1)
doc.add_paragraph(
    '商业在轨维护与服务是太空经济从"一次性"向"可持续"转型的关键基础设施。'
    '全球市场将从2025年的47亿美元增长至2035年的126亿美元，年复合增长率约10%。'
)
doc.add_paragraph(
    '国际层面，美国企业（诺格、Maxar）和日本企业（Astroscale）处于领先地位，'
    '2026年将是关键验证节点——美国太空军将实施4个重要任务演示。'
    '中国在GEO在轨加油技术已获世界领先成就，但商业化程度仍低。'
)
doc.add_paragraph(
    '对于投资者而言，Astroscale是全球唯一的纯正OOS上市标的；'
    '诺格是大型防务企业中最直接的受益者；'
    '国内投资则需等待中国商业航天进一步开放。'
)
doc.add_paragraph(
    '长期看，随着太空竞争加剧、卫星星座规模化运营，'
    '在轨服务将从"锦上添花"变为"不可或缺"，是值得战略性布局的赛道。'
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('免责声明：').bold = True
p.add_run('本报告仅供参考，不构成投资建议。投资有风险，决策需谨慎。')

# 保存
out = '/Users/mac/.openclaw/workspace/research/商业在轨维护与服务研究报告.docx'
doc.save(out)
print(f'✅ 报告已保存: {out}')
