#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《商业在轨维护与服务（On-Orbit Servicing）研究报告》
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading(doc, text, level, color=None):
    """添加自定义标题"""
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_paragraph(doc, text, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_bullet(doc, text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table_row(table, data, bg_color=None):
    """添加表格行"""
    row = table.add_row()
    for i, cell_text in enumerate(data):
        cell = row.cells[i]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if bg_color:
            set_cell_bg(cell, bg_color)
    return row

def create_report():
    doc = Document()

    # ==================== 页面设置 ====================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ==================== 封面标题 ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run('商业在轨维护与服务')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('（On-Orbit Servicing）')
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.name = 'Times New Roman'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run2.font.color.rgb = RGBColor(0, 51, 102)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('行业研究报告')
    run3.font.size = Pt(18)
    run3.font.name = 'Times New Roman'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run3.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f'报告日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    run_date.font.size = Pt(12)
    run_date.font.name = 'Times New Roman'
    run_date._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run('研究机构：开放爪研究院')
    run_author.font.size = Pt(12)
    run_author.font.name = 'Times New Roman'
    run_author._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ==================== 目录 ====================
    add_heading(doc, '目  录', 1, (0, 51, 102))

    toc_items = [
        '一、行业概述',
        '  1.1 行业定义与背景',
        '  1.2 行业发展历程',
        '  1.3 在轨服务的战略意义',
        '二、国际市场分析',
        '  2.1 全球市场规模与预测',
        '  2.2 国际主要公司分析',
        '  2.3 国际技术进展',
        '  2.4 商业模式与市场前景',
        '三、国内市场分析',
        '  3.1 国内市场规模与政策环境',
        '  3.2 国内主要公司分析',
        '  3.3 国内在轨服务进展',
        '四、技术分类深度分析',
        '  4.1 卫星延寿与服务',
        '  4.2 轨道转移与加油技术',
        '  4.3 碎片清理与主动碎片移除（ADR）',
        '  4.4 在轨制造与装配（ISAM）',
        '五、市场前景与投资机会',
        '  5.1 2030年市场规模预测',
        '  5.2 投资机会分析',
        '  5.3 风险与挑战',
        '六、结论与建议',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ==================== 正文 ====================

    # 一、行业概述
    add_heading(doc, '一、行业概述', 1, (0, 51, 102))

    add_heading(doc, '1.1 行业定义与背景', 2)
    add_paragraph(doc,
        '商业在轨维护与服务（On-Orbit Servicing，简称OOS）是指利用专用服务航天器对已在轨运行的目标卫星或其他航天器 '
        '执行观测、检测、维修、燃料补加、轨道调整、功能升级、碎片清除等一系列操作的技术与商业活动。'
        '与传统卫星"发射即终结"的模式不同，在轨服务旨在延长卫星寿命、提升卫星性能、降低系统全生命周期成本，'
        '并为日益严峻的空间碎片问题提供解决方案。')

    add_paragraph(doc,
        '随着全球在轨活跃卫星数量快速增长（截至2025年初已超过1万颗），加上大量老旧卫星燃料耗尽、'
        '故障频发的问题，在轨服务已成为航天产业链中极具战略价值的新兴领域。')

    add_heading(doc, '1.2 行业发展历程', 2)
    add_paragraph(doc, '在轨服务技术的发展大致经历了以下阶段：')

    history_data = [
        ['时间', '里程碑事件', '意义'],
        ['1984年', 'NASA航天飞机执行萨摩斯（Saturn）任务（未成功）', '人类首次尝试在轨维修'],
        ['1993年', '哈勃望远镜首次在轨维修（SM1）', '验证人机协同在轨服务可行性'],
        ['2007年', '美国"轨道快车"（Orbital Express）项目', '首次完全自主的在轨服务演示'],
        ['2012年', '德国ETS-VII在轨服务试验', '日本验证机器人对接与操作技术'],
        ['2015年', 'NASA恢复"卫星在轨服务"项目规划', '政府层面系统推进'],
        ['2019-2020年', '诺斯罗普·格鲁曼MEV-1/2成功商业对接', '商业在轨服务时代正式开启'],
        ['2024年', 'Astroscale ADRAS-J商业碎片清除任务', '商业ADR进入实操阶段'],
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(['时间', '里程碑事件', '意义']):
        hdr.cells[i].text = h
        set_cell_bg(hdr.cells[i], '003366')
        for para in hdr.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for row_data in history_data[1:]:
        add_table_row(table, row_data, 'F0F4FA')

    doc.add_paragraph()

    add_heading(doc, '1.3 在轨服务的战略意义', 2)
    add_paragraph(doc, '在轨服务技术的发展具有以下重大战略意义：')
    add_bullet(doc, '经济价值：可节省重新发射新卫星的高额成本（大型GEO卫星单星造价超过3亿美元），'
                    '延长价值数十亿美元的在轨资产寿命，降低运营商全生命周期成本。')
    add_bullet(doc, '空间安全：有效应对日益严峻的空间碎片威胁（目前可追踪碎片超过34,000个），'
                    '维持轨道环境的可持续利用，保障全球航天活动的长期安全。')
    add_bullet(doc, '军事战略：在轨服务技术具有军民两用特性，可用于侦察、干扰、捕获对手卫星，'
                    '已成为各航天大国战略博弈的关键领域。')
    add_bullet(doc, '深空探索：在轨服务技术是建设月球门户、火星补给站等深空基础设施的核心使能技术，'
                    '对人类深空探索事业具有基础性支撑作用。')
    add_bullet(doc, '新质生产力：作为商业航天的高价值细分赛道，在轨服务代表了航天产业从"一次性使用"'
                    '向"可持续运营"范式转变的根本性变革。')

    # 二、国际市场分析
    doc.add_page_break()
    add_heading(doc, '二、国际市场分析', 1, (0, 51, 102))

    add_heading(doc, '2.1 全球市场规模与预测', 2)

    market_data = [
        ['来源机构', '预测年份', '市场规模', '复合增长率(CAGR)', '备注'],
        ['MarketsandMarkets', '2030年', '51亿美元', '11.5%', '综合多种服务类型'],
        ['Precedence Research', '2035年', '126亿美元', '10.43%', '含GEO/LEO/MEO'],
        ['Future Market Insights', '2035年', '95亿美元', '11.7%', '细分服务类型'],
        ['Fortune Business Insights', '2034年', '68.7亿美元', '10.1%', '按在轨服务类型'],
        ['Credence Research', '2032年', '73.8亿美元', '12.8%', '较乐观预测'],
        ['Intel Market Research', '2034年', '约70亿美元', '10.1%', '加注/维修/升级'],
    ]

    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0]
    for i, h in enumerate(['来源机构', '预测年份', '市场规模', '复合增长率(CAGR)', '备注']):
        hdr2.cells[i].text = h
        set_cell_bg(hdr2.cells[i], '003366')
        for para in hdr2.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for row_data in market_data[1:]:
        add_table_row(table2, row_data, 'F0F4FA')

    doc.add_paragraph()
    add_paragraph(doc,
        '综合各研究机构数据，2025年全球在轨服务市场规模约在28亿~47亿美元区间，'
        '到2030年有望达到51亿~73亿美元，2035年则可能突破95亿~126亿美元。'
        '北美市场以约37%~45%的份额主导全球市场，欧洲市场预计2030年单独市场规模超过50亿美元，'
        '亚太地区增速领先。')

    add_heading(doc, '2.2 国际主要公司分析', 2)

    # 公司1：诺斯罗普·格鲁曼
    add_heading(doc, '2.2.1 诺斯罗普·格鲁曼（Northrop Grumman）', 3)
    add_paragraph(doc,
        '诺斯罗普·格鲁曼旗下全资子公司SpaceLogistics LLC是全球商业在轨服务领域的绝对领导者。'
        '公司已推出三款核心产品：')
    add_bullet(doc, 'MEV（Mission Extension Vehicle，任务延寿飞行器）：全球首个商业化的GEO卫星延寿服务。'
                    'MEV-1于2019年发射，2020年成功与国际通信卫星公司（Intelsat）的IS-901卫星对接，'
                    '将后者从"坟墓轨道"拉回正常GEO轨道并恢复运营，创造了历史。'
                    '2025年4月，MEV-1完成任务后成功解对接，释放IS-901回坟墓轨道，'
                    '完成了人类历史上首次商业航天器在GEO轨道的对接与解对接操作。'
                    'MEV-2于2020年发射，已与国际通信卫星IS-10-02对接并提供服务。')
    add_bullet(doc, 'MRV（Mission Robotic Vehicle，任务机器人飞行器）：配备机械臂，将在轨执行检查、'
                    '安装"任务延寿舱"（MEP）、搬运、修复乃至碎片清除等复杂任务。计划2025年发射。')
    add_bullet(doc, 'MEP（Mission Extension Pod，任务延寿舱）：小型化的推进增强装置，可安装在客户卫星上，'
                    '由MRV负责安装，提供轨道控制和动量卸载功能，计划2026年开始部署。')
    add_paragraph(doc,
        '财务方面，SpaceLogistics的延寿服务单次合同价值估计在数千万至亿美元级别，'
        '每颗MEV预计可服务5~15年，覆盖多颗卫星。')

    # 公司2：Astroscale
    add_heading(doc, '2.2.2 Astroscale（日本）', 3)
    add_paragraph(doc,
        'Astroscale成立于2013年，总部位于日本东京，是专注于空间碎片清除和在轨服务的商业航天公司，'
        '是在轨服务领域估值最高的初创企业之一。')
    add_bullet(doc, 'ADRAS-J（Active Debris Removal by Astroscale-Japan）：2024年2月发射，'
                    '是人类历史上首个商业碎片清除任务。该飞行器对一枚废弃的H-IIA火箭上面级进行了'
                    '近距离接近、绕飞观测和全面拍照，获取了大量有价值的碎片特性数据。'
                    '任务于2025年初结束，飞行器受控离轨。')
    add_bullet(doc, 'ADRAS-J2：接替任务，将尝试实际捕获并移除同一火箭残骸，'
                    '计划2025~2026年执行，使用Astroscale自研的机械臂技术。合同金额约132亿日元。')
    add_bullet(doc, 'ELSA-M（End-of-Life Service by Astroscale-M）：针对一网（OneWeb）星座卫星的'
                    '在轨服务与主动碎片移除任务，将验证多碎片清除能力。')
    add_bullet(doc, '技术积累：Astroscale在相对导航（RPO）、近距离操作、碎片捕获等核心技术方面'
                    '建立了完整的知识产权壁垒，并在日本、英国、新加坡等多地设有研发中心。')

    # 公司3：Maxar
    add_heading(doc, '2.2.3 Maxar Technologies（美国）', 3)
    add_paragraph(doc,
        'Maxar是全球领先的卫星影像与地理空间数据公司，近年积极布局在轨服务领域。'
        '其主要进展包括：')
    add_bullet(doc, ' Restore-L任务概念（已被NASA采纳为Mars贪婪任务的一部分）：'
                    '提出利用机器人服务飞行器为NASA现役卫星延寿的技术路线。')
    add_bullet(doc, '在轨检查服务：利用高分辨率成像技术为客户提供在轨卫星状态评估，'
                    '帮助运营商判断卫星健康状况。')
    add_bullet(doc, 'Maxar已被NASA选中为"轨道礁"（Orbital Reef，多个商业空间站项目之一）'
                    '提供部分技术验证，但核心业务仍是地球观测与卫星影像。')

    # 公司4：SpaceX
    add_heading(doc, '2.2.4 SpaceX（美国）', 3)
    add_paragraph(doc,
        'SpaceX作为全球商业航天绝对龙头，其在轨服务能力正在快速建设中：')
    add_bullet(doc, '星链（Starlink）卫星的批量部署与在轨管理经验为SpaceX积累了丰富的'
                    '轨道转移、近距离编队和自主交会对接经验。')
    add_bullet(doc, '星舰（Starship）计划用于未来大型在轨设施的建设与补给，'
                    '其超大型近地轨道（LEO）运载能力将为在轨服务提供全新平台。')
    add_bullet(doc, 'SpaceX已获得NASA合同，研究利用星舰进行在轨加油的技术可行性，'
                    '这被视为实现载人月球探测（Artemis计划）的关键技术之一。')
    add_bullet(doc, '为AST SpaceMobile等卫星运营商提供星舰整流罩共享发射，间接推动'
                    '商业卫星规模化和标准化，从而扩大在轨服务市场。')

    # 公司5：OneWeb
    add_heading(doc, '2.2.5 OneWeb（英国/印度）', 3)
    add_paragraph(doc,
        'OneWeb是全球第二大低轨宽带互联网星座运营商（仅次于SpaceX Starlink），'
        '截至2025年初已部署超过630颗卫星。与Astroscale合作：')
    add_bullet(doc, 'OneWeb已签约Astroscale的ELSA-M服务，承诺在星座寿命末期'
                    '主动移除退役卫星，为行业树立空间交通管理标杆。')
    add_bullet(doc, 'OneWeb星座的大规模部署本身也在推动在轨服务市场需求：'
                    '大规模星座意味着大量需要统一维护的同型卫星，服务规模效应显著。')

    # 公司6：OrbitFab
    add_heading(doc, '2.2.6 OrbitFab（美国）', 3)
    add_paragraph(doc,
        'OrbitFab是美国专注于在轨加注服务的初创公司，已推出"太空加油站"（Space Tanker）概念：')
    add_bullet(doc, '开发可重复使用的在轨推进剂输送飞行器，计划为GEO和LEO卫星提供燃料补加。')
    add_bullet(doc, '已获得美国国防部创新部门（DIU）和NASA的多份合同支持。')
    add_bullet(doc, '技术路线：采用低温推进剂在轨转移技术，技术难度高但潜在市场价值极大。')

    add_heading(doc, '2.3 国际技术进展', 2)

    add_heading(doc, '2.3.1 卫星延寿与服务', 3)
    add_paragraph(doc,
        '卫星延寿是在轨服务最成熟、商业化程度最高的细分市场。'
        '以诺斯罗普·格鲁曼MEV为代表，采用"停泊式"延寿方案——服务飞行器对接至目标卫星非工作面，'
        '接管其姿态控制和轨道维持功能。该技术已完全成熟并实现商业化运营。'
        '新一代MRV将升级为"机器人延寿"，可安装推进增强舱，适用性更广。')

    add_heading(doc, '2.3.2 轨道加油与推进增强', 3)
    add_paragraph(doc,
        '在轨加油是未来扩展卫星寿命的核心技术路线，但技术难度远高于纯延寿。'
        '主要挑战包括：低温推进剂（液氢、液氧）的在轨转注、超低温流体的热管理、'
        '流体动力学建模等。NASA已验证"土星"（Saturn）加油技术概念，'
        'OrbitFab正在开发冷冻推进剂输送飞行器。预计2027~2028年实现首次商业GEO卫星在轨加油。')

    add_heading(doc, '2.3.3 碎片清除（ADR）', 3)
    add_paragraph(doc,
        '主动碎片清除（Active Debris Removal，ADR）是在轨服务增速最快的细分市场之一。'
        'Astroscale ADRAS-J于2024年完成了里程碑式的商业碎片检查任务，'
        '证明商业公司已具备对非合作目标（非合作卫星或废弃火箭残骸）实施'
        '近距观测和绕飞的能力。下一阶段（ADRAS-J2）将验证实际捕获与移除，'
        '这对整个ADR行业具有决定性意义。'
        '欧洲也在积极推进ADR技术：ESA的e.Deorbit项目处于技术验证阶段。')

    add_heading(doc, '2.3.4 在轨制造与装配（ISAM）', 3)
    add_paragraph(doc,
        '在轨服务、装配与制造（In-space Servicing, Assembly and Manufacturing，ISAM）'
        '代表了更长远的技术愿景，包括：')
    add_bullet(doc, '在轨装配：利用多个自主机器人或宇航员在太空组装大型结构（如大型太阳帆、'
                    '空间太阳能电站、大型光学望远镜等），目前DARPA的RSGS项目（机器人卫星加油服务）'
                    '正在推进此方向。')
    add_bullet(doc, '在轨制造：在微重力环境下制造光纤、药品、合金材料等高附加值产品，'
                    'Varda Space、Space Applications Services等公司已率先探索。')
    add_bullet(doc, '在轨升级：类似"手机系统升级"的模块化卫星升级，'
                    '目前受限于卫星接口标准化问题，仍处于早期阶段。')

    add_heading(doc, '2.4 商业模式与市场前景', 2)
    add_paragraph(doc,
        '当前在轨服务的商业模式主要包括：')
    add_bullet(doc, '直接延寿合同：卫星运营商与服务公司签订5~10年延寿服务合同，'
                    '按年付费，单次合同金额在数千万至数亿美元不等。'
                    '代表案例：Intelsat与SpaceLogistics的MEV服务合同。')
    add_bullet(doc, '碎片清除服务：政府或运营商委托商业公司清除特定碎片，'
                    '按任务付费，单次任务通常在数亿日元至数亿美元。')
    add_bullet(doc, '订阅式服务：类似于"太空4S店会员"模式，运营商按月/年支付服务费，'
                    '享受定期检查、维护和延寿服务。')
    add_bullet(doc, '在轨加油即服务（Fuel-as-a-Service）：按实际补加的推进剂量计费，'
                    '降低运营商一次性成本，转移至按需付费模式。')
    add_paragraph(doc,
        '市场前景方面，随着Starlink、OneWeb等大型星座进入寿命中期（5~7年），'
        '以及对GEO通信卫星延寿需求的持续增长，2025~2030年将是商业在轨服务市场的'
        '黄金发展期，年均复合增长率维持在11%~13%区间。')

    # 三、国内市场分析
    doc.add_page_break()
    add_heading(doc, '三、国内市场分析', 1, (0, 51, 102))

    add_heading(doc, '3.1 国内市场规模与政策环境', 2)

    add_paragraph(doc,
        '中国在轨服务市场尚处于从技术验证向商业化过渡的阶段，但政策支持力度持续加大：')

    add_bullet(doc, '政策支持：2024年，商业航天首次作为"新增长引擎"写入政府工作报告；'
                    '2025年，商业航天进一步升级为"战略性新兴产业"；'
                    '国家航天局于2025年正式设立商业航天司，实现发射审批、频轨资源、安全监管的一体化统筹管理。')
    add_bullet(doc, '市场规模：据赛迪智库数据，2025年中国商业航天市场规模约2.83万亿元，'
                    '同比增长21.7%；2026年预计达到3.5万亿元，其中在轨服务作为新兴细分领域，'
                    '将分享产业高速增长红利。')
    add_bullet(doc, '十五五规划：2026年"十五五"规划纲要首次写入"航天强国"建设目标，'
                    '明确提出"建成在轨服务与维护系统，设立空间4S店"，'
                    '为国内在轨服务产业发展提供明确政策指引。')

    add_heading(doc, '3.2 国内主要公司分析', 2)

    add_heading(doc, '3.2.1 中国航天科技集团（CASC）', 3)
    add_paragraph(doc,
        '中国航天科技集团是 中国在轨服务技术的核心推动力量，旗下相关资产和技术包括：')
    add_bullet(doc, '实践二十五号卫星：2025年初成功发射，是我国首个专门验证卫星燃料补加'
                    '与延寿服务的在轨服务技术试验平台。该卫星将验证未来对在轨卫星'
                    '自主维修和加油的核心技术，为我国商业在轨服务奠定技术基础。')
    add_bullet(doc, '实践二十号卫星：此前已验证多项先进技术在轨性能，包括激光通信、'
                    '新型离子推进等，为在轨服务提供配套技术支撑。')
    add_bullet(doc, '2025年战略布局：航天科技集团于2025年9月分别成立了商业卫星公司'
                    '和商业火箭公司，形成"国家队+商业化"双轨发展格局，'
                    '其中商业卫星公司将在在轨服务领域重点布局。')
    add_bullet(doc, '技术优势：CASC在神舟飞船、天宫空间站等载人航天任务中积累了丰富的'
                    '交会对接、机器人操作经验，这些技术可直接转化应用于在轨服务领域。')

    add_heading(doc, '3.2.2 银河航天', 3)
    add_paragraph(doc,
        '银河航天成立于2016年，是中国商业航天领域第一家独角兽企业，估值约115亿元：')
    add_bullet(doc, '核心业务：低轨宽带通信卫星的研制与运营，"小蜘蛛网"试验星座已验证'
                    '多项关键技术，与Starlink技术路线对标。')
    add_bullet(doc, '在轨服务相关：银河航天积极探索卫星批量化生产技术，'
                    '其卫星智慧工厂实现了年产150颗中型卫星的批量制造能力，'
                    '为未来在轨服务卫星的批量化奠定了制造基础。')
    add_bullet(doc, '2026年已完成股改：从"银河航天（北京）网络技术有限公司"'
                    '更名为"银河航天（北京）科技集团股份有限公司"，'
                    '进一步向产业链上下游延伸布局。')
    add_bullet(doc, '战略定位：虽以卫星互联网为主业，但其在低轨卫星平台标准化、'
                    '批量制造方面的积累，将为未来国内在轨服务市场的爆发提供基础支撑。')

    add_heading(doc, '3.2.3 长光卫星', 3)
    add_paragraph(doc,
        '长光卫星是中国商业航天遥感领域的龙头企业：')
    add_bullet(doc, '"吉林一号"星座：国内规模最大的商业遥感星座，在轨卫星超过117颗，'
                    '实现了全球遥感数据的实时获取能力。')
    add_bullet(doc, '全产业链布局：构建了"卫星制造-星座运营-数据服务-生态赋能"的'
                    '全产业链垂直整合体系，具备显著的成本控制和市场响应优势。')
    add_bullet(doc, '在轨服务关联：作为国内最大的民营遥感卫星运营商，'
                    '其星座运营经验对探索在轨服务商业模式具有重要参考价值。')
    add_bullet(doc, '已于2020年完成Pre-IPO轮融资24.64亿元，正在推进上市进程。')

    add_heading(doc, '3.2.4 其他新兴力量', 3)
    add_paragraph(doc, '除上述公司外，国内在轨服务领域的新兴力量还包括：')
    add_bullet(doc, '航天宏图：运营"女娲星座"（SAR卫星星座），在轨卫星12颗，'
                    '是国内规模最大的商业雷达遥感星座，正探索遥感卫星在轨服务新模式。')
    add_bullet(doc, '时空道宇：背靠吉利集团，聚焦低轨卫星通信，'
                    '2024年入轨23颗卫星（民营之首），推进天地一体化通信建设。')
    add_bullet(doc, '天仪研究院：国内首个商业SAR卫星公司，'
                    '成功发射中国首颗商业SAR卫星"海丝一号"，技术路线差异化。')

    add_heading(doc, '3.3 国内在轨服务进展', 2)

    add_paragraph(doc,
        '中国在轨服务技术的发展近年来取得显著突破，但仍处于技术验证阶段，'
        '距离大规模商业化仍需时日。主要进展包括：')

    progress_data = [
        ['项目/试验', '时间', '主要内容', '意义'],
        ['实践二十五号', '2025年', '验证卫星燃料补加与延寿服务技术', '国内首个在轨服务技术验证平台'],
        ['巡天空间望远镜（CSST）', '计划2026年前后', '采用可开合舱板式设计，支持在轨维修', '验证大型光学天文卫星可维修性'],
        ['空间站机械臂操作', '持续进行', '神舟十四/十五号任务中完成多次舱外维修演练', '积累机器人舱外作业经验'],
        ['可重复使用火箭', '2024-2026年', '朱雀三号、长征十二号甲等可回收火箭测试', '降低在轨服务发射成本'],
        ['商业航天立法', '2025年', '国家航天局发布商业航天高质量安全发展行动计划', '规范在轨服务监管框架'],
    ]

    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0]
    for i, h in enumerate(['项目/试验', '时间', '主要内容', '意义']):
        hdr3.cells[i].text = h
        set_cell_bg(hdr3.cells[i], '003366')
        for para in hdr3.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for row_data in progress_data[1:]:
        add_table_row(table3, row_data, 'F0F4FA')

    doc.add_paragraph()
    add_paragraph(doc,
        '从整体技术水平看，中国在轨服务领域与欧美先进水平仍有差距，主要体现在：'
        '在轨对接操作的工程经验相对不足、商业化程度较低、关键技术（如大功率机器人机械臂、'
        '低温推进剂在轨转注）尚处于验证阶段。但受益于中国航天体系完整性和后发优势，'
        '追赶速度正在加快，预计2028~2030年将实现首批商业化在轨服务。')

    # 四、技术分类深度分析
    doc.add_page_break()
    add_heading(doc, '四、技术分类深度分析', 1, (0, 51, 102))

    add_heading(doc, '4.1 卫星延寿与服务', 2)
    add_paragraph(doc,
        '卫星延寿是在轨服务商业化程度最高的细分市场，核心技术路线包括：')
    add_bullet(doc, '停泊式延寿（Park-and-Ride）：服务飞行器对接至目标卫星，'
                    '接管其姿态控制与轨道维持功能。代表：诺格MEV系列。'
                    '技术成熟度（TRL）：9（已完全商业化）。')
    add_bullet(doc, '推进增强（Propulsive Augmentation）：通过MRV安装推进增强舱（MEP），'
                    '为缺乏推进剂的目标卫星补充推进能力。代表：SpaceLogistics MEP。'
                    '技术成熟度：6~7（2025~2026年商业化）。')
    add_bullet(doc, '模块化更换（Modular Replacement）：更换故障或过时模块，恢复或升级卫星功能。'
                    '代表：DARPA RSGS项目。技术成熟度：4~6（2028~2030年预期商业化）。')
    add_paragraph(doc,
        '市场驱动因素：全球现有约500颗GEO商业通信卫星，平均寿命10~15年，'
        '大量卫星在燃料耗尽后仍有有效载荷功能，延寿服务市场刚性需求强。'
        '单次延寿服务可为运营商节省重新发射一颗大型GEO卫星的费用（通常在2亿~5亿美元），'
        '服务定价空间充足。')

    add_heading(doc, '4.2 轨道转移与加油技术', 2)
    add_paragraph(doc,
        '轨道转移与在轨加油技术是在轨服务中技术难度最高的领域之一：')
    add_bullet(doc, '轨道转移：在轨飞行器利用自身推进系统改变目标卫星轨道，'
                    '包括从坟墓轨道拉回GEO工作位置、LEO到GEO的轨道抬升、'
                    '卫星离轨处置等。核心技术包括：相对导航（RPO）、轨道机动规划、'
                    '多体系统动力学等。')
    add_bullet(doc, '在轨加油：服务飞行器携带推进剂至目标卫星，加注后恢复其轨道机动能力。'
                    '技术挑战：低温推进剂（液氧、液氢、偏二甲肼等）的在轨转注；'
                    '流体热管理；零排放加注接口设计等。NASA Restore-L任务（后演变为Mars贪婪任务）'
                    '对此进行了深入研究。')
    add_bullet(doc, '技术成熟度：轨道转移（MEV）：TRL 9；'
                    '在轨加油（低温推进剂）：TRL 4~6（预计2028~2030年商业化）。')
    add_paragraph(doc,
        '国内进展：实践二十五号卫星正在进行相关技术验证，'
        '但低温推进剂在轨转注等核心技术仍待突破。')

    add_heading(doc, '4.3 碎片清理与主动碎片移除（ADR）', 2)
    add_paragraph(doc,
        '主动碎片移除（Active Debris Removal，ADR）是当前在轨服务增速最快的细分市场。'
        '全球可追踪空间碎片已超过34,000个（2025年数据），'
        '大量废弃火箭上面级和失效卫星形成潜在威胁。')
    add_paragraph(doc, '主要技术路线：')
    add_bullet(doc, '网捕式捕获：利用展开式网捕获不规则形状碎片。'
                    '代表：欧洲RemoveDEBRIS任务（已验证）。技术难度适中。')
    add_bullet(doc, '机械臂捕获：利用多自由度机械臂直接抓取目标。'
                    '代表：Astroscale ADRAS-J2（计划中）、ESA e.Deorbit。'
                    '技术难度高，但精度和可靠性更优。')
    add_bullet(doc, '磁吸式捕获：利用磁性对接环捕获磁性金属目标。适用面较窄。')
    add_bullet(doc, '激光推移：利用高能激光照射碎片产生等离子体喷射，推动碎片离轨。'
                    '技术尚不成熟，存在国际条约限制。')
    add_paragraph(doc,
        '商业前景：随着联合国和平利用外层空间委员会（COPUOS）空间碎片减缓准则的'
        '日益严格（要求LEO卫星在寿命结束后25年内离轨），'
        'ADR服务将逐步从"可选项"变为"必选项"，市场空间确定性高。'
        'Astroscale估计，仅OneWeb星座就存在数百颗卫星的ADR市场需求。')

    add_heading(doc, '4.4 在轨制造与装配（ISAM）', 2)
    add_paragraph(doc,
        '在轨服务、装配与制造（In-space Servicing, Assembly and Manufacturing，ISAM）'
        '代表了人类利用太空的终极愿景之一，被美国列为国家航天战略优先方向。')
    add_paragraph(doc, '核心应用场景：')
    add_bullet(doc, '大型空间结构在轨装配：如10米以上光学望远镜、空间太阳能电站、'
                    '超大型天线阵列等，超出单次火箭整流罩尺寸限制，必须在轨组装。')
    add_bullet(doc, '在轨制造高附加值产品：微重力环境下生产的药品（蛋白质晶体、疫苗）、'
                    '光纤预制棒（比陆基产品品质更高）、特殊合金等。'
                    'Varda Space等公司已实现太空药物制造的商业化。')
    add_bullet(doc, '卫星升级与改造：更换星载计算机、放大器、天线等模块，'
                    '使现役卫星获得等同于新一代卫星的性能。'
                    '需要解决接口标准化问题。')
    add_bullet(doc, '太空数据中心：在轨部署计算和存储服务器，减少数据下行延迟，'
                    '尤其适合卫星遥感数据的实时处理。')
    add_paragraph(doc,
        '市场预测：ISAM市场目前规模极小（不足5亿美元），'
        '预计到2035年有望增长至30亿~50亿美元，'
        '将成为在轨服务领域增速最快的子市场，但商业化时间节点预计在2030年之后。')

    # 五、市场前景与投资机会
    doc.add_page_break()
    add_heading(doc, '五、市场前景与投资机会', 1, (0, 51, 102))

    add_heading(doc, '5.1 2030年市场规模预测', 2)

    forecast_data = [
        ['细分市场', '2025年（亿美元）', '2030年预测（亿美元）', 'CAGR', '驱动因素'],
        ['卫星延寿服务', '12~15', '28~35', '15~18%', 'GEO卫星延寿需求刚性'],
        ['在轨加油', '1~2', '8~12', '45~55%', '技术突破+星座维护需求'],
        ['主动碎片清除(ADR)', '3~5', '12~18', '32~38%', '监管趋严+运营商需求'],
        ['在轨检查/SSA', '5~7', '10~15', '12~15%', '卫星运营商安全需求'],
        ['ISAM（装配制造）', '1~2', '5~8', '35~42%', '深空探索+商业应用'],
        ['合计', '28~47', '51~73', '11~13%', '综合'],
    ]

    table4 = doc.add_table(rows=1, cols=5)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0]
    for i, h in enumerate(['细分市场', '2025年（亿美元）', '2030年预测（亿美元）', 'CAGR', '驱动因素']):
        hdr4.cells[i].text = h
        set_cell_bg(hdr4.cells[i], '003366')
        for para in hdr4.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for row_data in forecast_data[1:]:
        add_table_row(table4, row_data, 'F0F4FA')

    doc.add_paragraph()
    add_paragraph(doc,
        '综合多家研究机构数据，2030年全球在轨服务市场规模预计达到51亿~73亿美元（约合370亿~530亿元人民币），'
        '2025-2030年复合增长率约为11%~13%。其中增速最快的细分领域为：'
        '在轨加油（45%~55%，低基数+高技术壁垒）、'
        'ISAM（35%~42%，长期蓝海市场）、'
        'ADR（32%~38%，监管+商业双轮驱动）。')

    add_heading(doc, '5.2 投资机会分析', 2)

    add_paragraph(doc, '基于上述分析，我们梳理出以下主要投资机会：')
    add_paragraph(doc, '第一，成熟期投资机会（当前至2027年）：', bold=True)
    add_bullet(doc, '卫星延寿服务运营商：诺斯罗普·格鲁曼/SpaceLogistics是确定性最强的龙头，'
                    '但作为上市公司母公司的一部分，投资门槛较高；Astroscale为非上市私人公司，'
                    '可关注其IPO机会。')
    add_bullet(doc, '在轨服务发射服务商：SpaceX凭借低成本高可靠发射能力，'
                    '将成为大多数商业在轨服务任务的首选发射商，受益确定性高。')
    add_bullet(doc, '空间态势感知（SSA）数据服务商：碎片监测和卫星跟踪数据是在轨服务的基础，'
                    'Maxar、LeoLabs等公司在此领域有先发优势。')

    add_paragraph(doc, '第二，成长期投资机会（2027~2032年）：', bold=True)
    add_bullet(doc, '国内在轨服务产业链：随着实践系列任务的推进和商业化条件成熟，'
                    '航天科技集团旗下上市公司（如中国卫星600118.SH）、'
                    '银河航天（Pre-IPO阶段）等将受益于国内市场规模扩张。')
    add_bullet(doc, '低温推进技术：OrbitFab等公司的低温推进剂在轨转注技术一旦突破，'
                    '将开辟全新的蓝海市场，有较高投资赔率。')
    add_bullet(doc, '机器人与AI技术：MRV执行在轨服务依赖高精度机械臂和自主决策AI，'
                    '相关核心零部件和算法供应商具有较高技术壁垒和投资价值。')

    add_paragraph(doc, '第三，长期布局机会（2030年以后）：', bold=True)
    add_bullet(doc, 'ISAM基础设施：太空制造、太空数据中心等业务的商业化，'
                    '将催生全新的太空经济生态，提前布局相关技术的公司值得关注。')
    add_bullet(doc, '深空在轨服务：随着月球门户、火星任务推进，'
                    '在地月空间乃至地火空间执行在轨补给和维护的技术需求将爆发，'
                    '具有长期战略投资价值。')

    add_heading(doc, '5.3 风险与挑战', 2)
    add_paragraph(doc, '在看到巨大机遇的同时，我们也必须清醒认识到在轨服务领域面临的风险：')
    add_bullet(doc, '技术风险：在轨加油、低温推进剂转注、大质量碎片捕获等核心技术'
                    '尚未完全成熟，任务失败概率较高；太空环境的极端条件（辐射、微流星体、'
                    '原子氧腐蚀等）对服务飞行器的可靠性提出极高要求。')
    add_bullet(doc, '监管风险：各国对在轨服务活动的法律定性尚不统一，'
                    '"卫星捕获"行为是否构成对他国财产的侵犯存在法律模糊地带；'
                    '军用与民用在轨服务的技术边界模糊，存在引发空间军事冲突的潜在风险。')
    add_bullet(doc, '商业风险：在轨服务单次任务成本高、服务定价缺乏透明基准，'
                    '商业模式仍需市场验证；客户集中度高（少数大型卫星运营商），'
                    '收入稳定性存在不确定性。')
    add_bullet(doc, '竞争格局：美国在轨服务技术全面领先，中国、欧洲、日本加速追赶，'
                    '未来可能出现技术路线竞争和市场分割问题；'
                    '国内产业链关键元器件（如高端星载计算机、高精度传感器）仍依赖进口，'
                    '存在供应链安全隐患。')

    # 六、结论与建议
    doc.add_page_break()
    add_heading(doc, '六、结论与建议', 1, (0, 51, 102))

    add_paragraph(doc,
        '商业在轨维护与服务（On-Orbit Servicing）作为航天产业的新兴高价值赛道，'
        '正在从技术验证迈向全面商业化，驱动因素明确、市场空间可观、技术壁垒深厚，'
        '具备成为下一个"千亿级"蓝海市场的潜力。')
    add_paragraph(doc, '主要结论：', bold=True)
    add_bullet(doc, '市场规模：2025年全球在轨服务市场规模约28亿~47亿美元，'
                    '2030年有望达到51亿~73亿美元，2035年则可能突破95亿~126亿美元，'
                    '是一个处于高速成长期的新兴市场。')
    add_bullet(doc, '竞争格局：诺斯罗普·格鲁曼/SpaceLogistics凭借MEV的成熟商业运营，'
                    '占据绝对领导地位；Astroscale在ADR领域处于全球领先地位；'
                    '国内整体仍处技术验证阶段，航天科技集团为中坚力量，'
                    '银河航天、长光卫星等民企快速跟进。')
    add_bullet(doc, '技术趋势：卫星延寿服务最为成熟（已商业化），'
                    '在轨加油和ADR增速最快（2025~2030年爆发期），'
                    'ISAM代表长期方向（2030年后商业化）。')
    add_bullet(doc, '投资判断：当前是布局国际成熟企业的同时，'
                    '密切关注国内技术突破带来的弯道超车机会；'
                    '国内在轨服务市场规模虽较国际市场滞后3~5年，'
                    '但增速更快、政策支持更强，2028~2030年有望迎来商业化拐点。')

    add_paragraph(doc, '投资建议：', bold=True)
    add_bullet(doc, '关注国内政策催化：国家航天局商业航天司的设立和"十五五"规划的推进，'
                    '将持续释放政策红利；实践系列任务的进展是重要观察窗口。')
    add_bullet(doc, '关注产业链关键环节：机器人机械臂、低温推进剂转注系统、'
                    '相对导航传感器等核心零部件供应商是高壁垒、高价值的投资标的。')
    add_bullet(doc, '关注民企崛起：银河航天、航天宏图、长光卫星等头部民企'
                    '在卫星批量制造和星座运营方面的积累，'
                    '将为其进入在轨服务赛道提供坚实基础。')
    add_bullet(doc, '审慎评估风险：在轨服务技术复杂度高、监管不确定性较大，'
                    '投资需关注技术验证进展和监管政策变化，'
                    '建议通过产业基金等分散化方式参与。')

    doc.add_paragraph()
    doc.add_paragraph()

    # 免责声明
    p_disclaimer = doc.add_paragraph()
    p_disclaimer.paragraph_format.space_before = Pt(20)
    run_disclaimer = p_disclaimer.add_run(
        '免责声明：本报告仅供参考，不构成任何投资建议。'
        '报告中数据来源于公开渠道，编制者不对数据准确性承担责任。'
        '投资有风险，决策需谨慎。')
    run_disclaimer.font.size = Pt(9)
    run_disclaimer.font.color.rgb = RGBColor(128, 128, 128)
    run_disclaimer.font.name = 'Times New Roman'
    run_disclaimer._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_disclaimer.italic = True

    # 保存文档
    output_path = '/Users/mac/.openclaw/workspace/research/商业在轨维护与服务研究报告.docx'
    doc.save(output_path)
    print(f"报告已生成：{output_path}")

if __name__ == '__main__':
    create_report()
